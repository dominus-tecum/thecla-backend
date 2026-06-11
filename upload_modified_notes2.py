import os
import re
import uuid
import base64
from docx import Document
from docx.shared import Inches
import requests
import PyPDF2
from PIL import Image
import io
import re
# ============================
# FILE SORTING FUNCTIONS
# ============================

def extract_chapter_number(filename):
    """Extract chapter number from filename for sorting"""
    # Match patterns like "Chapter 1", "Chapter 10", "Chapter 2"
    match = re.search(r'Chapter\s+(\d+)', filename, re.IGNORECASE)
    if match:
        return int(match.group(1))
    
    # Also match patterns like "1. ", "01. ", "1-", etc.
    match = re.search(r'^(\d+)[\.\s-]', filename)
    if match:
        return int(match.group(1))
    
    return 999  # Put files without chapter number at the end

def sort_files_by_chapter(files):
    """Sort files by chapter number naturally"""
    return sorted(files, key=extract_chapter_number)

# ============================
# DEBUG FUNCTIONS
# ============================

def debug_pdf_content(pdf_path):
    """Debug what's actually in the PDF"""
    print(f"\n🔍 DEBUGGING PDF: {pdf_path}")
    
    # Check if file exists
    if not os.path.exists(pdf_path):
        print(f"❌ File not found: {pdf_path}")
        return False
    
    # Check file size
    file_size = os.path.getsize(pdf_path)
    print(f"   File size: {file_size} bytes ({file_size/1024:.1f} KB)")
    
    # Method 1: Try reading as binary to see if it's actually a PDF
    try:
        with open(pdf_path, 'rb') as f:
            first_bytes = f.read(20)  # Read more bytes for better detection
            print(f"   File signature (hex): {first_bytes.hex()[:40]}...")
            
            # Check for PDF signature
            if b'%PDF' in first_bytes:
                print("   ✓ Valid PDF signature found")
                pdf_version = first_bytes[5:8].decode('ascii', errors='ignore')
                print(f"   PDF version: {pdf_version}")
            else:
                print("   ✗ WARNING: Not a valid PDF file (missing %PDF header)")
                return False
    except Exception as e:
        print(f"   File read error: {e}")
        return False
    
    # Method 2: Try PyPDF2
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            print(f"\n   PyPDF2 Analysis:")
            print(f"   - Pages found: {len(reader.pages)}")
            print(f"   - Is encrypted: {reader.is_encrypted}")
            print(f"   - Metadata: {reader.metadata}")
            
            # Check text extraction on first 3 pages
            pages_with_text = 0
            total_text_chars = 0
            
            for i in range(min(3, len(reader.pages))):
                text = reader.pages[i].extract_text()
                text_length = len(text.strip())
                total_text_chars += text_length
                
                if text_length > 0:
                    pages_with_text += 1
                    print(f"\n   Page {i+1}:")
                    print(f"   - Has text: ✓ ({text_length} characters)")
                    print(f"   - First 300 chars: '{text[:300].replace(chr(10), ' ').replace(chr(13), ' ')}...'")
                    
                    # Look for specific patterns
                    if 'image[[' in text:
                        print(f"   - Contains image markers: ✓")
                    if '##' in text:
                        print(f"   - Contains markdown headers: ✓")
                else:
                    print(f"\n   Page {i+1}:")
                    print(f"   - Has text: ✗ (0 characters)")
                    print(f"   - WARNING: Page appears to be empty or image-based")
            
            print(f"\n   Summary:")
            print(f"   - Pages with text: {pages_with_text}/{min(3, len(reader.pages))}")
            print(f"   - Total text chars in first {min(3, len(reader.pages))} pages: {total_text_chars}")
            
            if total_text_chars == 0:
                print("   ⚠️  CRITICAL: No text extracted! PDF is likely scanned/image-based")
                return False
            elif total_text_chars < 100:
                print("   ⚠️  WARNING: Very little text extracted")
                return True
            else:
                print("   ✓ Good amount of text extracted")
                return True
                
    except Exception as e:
        print(f"   PyPDF2 error: {e}")
        import traceback
        traceback.print_exc()
        return False

def debug_docx_content(docx_path):
    """Debug what's in a Word document"""
    print(f"\n🔍 DEBUGGING DOCX: {docx_path}")
    
    try:
        document = Document(docx_path)
        
        # Count paragraphs
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        print(f"   Paragraphs with text: {len(paragraphs)}")
        
        # Show first few paragraphs
        for i, para in enumerate(paragraphs[:5]):
            print(f"   Para {i+1}: '{para[:100]}...'")
        
        # Check for images
        images_count = 0
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                images_count += 1
        print(f"   Images found: {images_count}")
        
        # Check for tables
        print(f"   Tables found: {len(document.tables)}")
        
        return True
        
    except Exception as e:
        print(f"   DOCX error: {e}")
        return False

# ============================
# IMAGE DIAGNOSTIC FUNCTIONS
# ============================

def check_debug_file_for_markers():
    """Check the saved debug file for image markers"""
    debug_file = "pdf_extracted_text_debug.txt"
    
    if os.path.exists(debug_file):
        print(f"\n🔍 CHECKING DEBUG FILE: {debug_file}")
        
        with open(debug_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
            # Simple check
            if 'image[[' in content:
                print("✅ 'image[[' FOUND in debug file!")
                
                # Count occurrences
                count = content.count('image[[')
                print(f"   Found {count} occurrences")
                
                # Show first few
                lines = content.split('\n')
                found_count = 0
                for i, line in enumerate(lines):
                    if 'image[[' in line:
                        found_count += 1
                        print(f"   Line {i+1}: {line.strip()[:100]}")
                        if found_count >= 3:  # Show only first 3
                            break
            else:
                print("❌ 'image[[' NOT FOUND in debug file")
                print("   This confirms PyPDF2 isn't extracting them")
                
                # Check for variations
                variations = [
                    r'image\s*\[\[',  # image [[ with spaces
                    r'image\[',       # image[ without second bracket
                    r'\[\[image',     # [[image reversed
                    r'IMG\[\[',       # uppercase
                    r'Image\[\[',     # capitalized
                ]
                
                import re
                for pattern in variations:
                    if re.search(pattern, content):
                        print(f"   ⚠️  Found variation: '{pattern}'")
                        
    else:
        print(f"❌ Debug file not found: {debug_file}")

def check_pdf_for_hidden_markers(pdf_path):
    """Check if image markers exist but aren't being extracted by PyPDF2"""
    print(f"\n🔍 DEEP CHECK for markers in: {os.path.basename(pdf_path)}")
    
    try:
        # 1. Check raw PDF bytes
        print("  1. Scanning raw PDF bytes...")
        with open(pdf_path, 'rb') as file:
            raw_content = file.read()
            content_str = raw_content.decode('latin-1', errors='ignore')
            
            # Look for variations
            variations = [
                ('image[[', 'Standard marker'),
                ('image [ [', 'Marker with spaces'),
                ('image[', 'Missing bracket'),
                ('[[image', 'Reversed'),
                ('IMG[[', 'Uppercase'),
                ('Image[[', 'Capitalized'),
            ]
            
            found_any = False
            for pattern, description in variations:
                count = content_str.count(pattern)
                if count > 0:
                    print(f"     ✅ Found '{pattern}' ({description}): {count} times")
                    found_any = True
            
            if not found_any:
                print("     ❌ No image markers found in raw PDF bytes")
                print("     The markers might be in images or vector graphics")
        
        # 2. Try PyMuPDF (if installed) - better text extraction
        print("\n  2. Trying alternative extraction methods...")
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(pdf_path)
            pymupdf_text = ""
            for page in doc:
                pymupdf_text += page.get_text()
            doc.close()
            
            if 'image[[' in pymupdf_text:
                count = pymupdf_text.count('image[[')
                print(f"     ✅ PyMuPDF found 'image[[': {count} times")
                # Show sample
                idx = pymupdf_text.find('image[[')
                if idx != -1:
                    sample = pymupdf_text[max(0, idx-50):idx+100]
                    print(f"     Sample: ...{sample}...")
            else:
                print("     ❌ PyMuPDF also didn't find 'image[['")
                
        except ImportError:
            print("     ℹ️  PyMuPDF not installed (pip install PyMuPDF)")
        except Exception as e:
            print(f"     ⚠️  PyMuPDF error: {e}")
            
        print("\n📊 CONCLUSION:")
        if 'image[[' in content_str:
            print("   The PDF contains image markers but PyPDF2 isn't extracting them.")
            print("   Consider using PyMuPDF for better text extraction.")
        else:
            print("   The PDF doesn't contain 'image[[' markers in any extractable form.")
            print("   The markers you see might be in images or special graphics.")
            
    except Exception as e:
        print(f"❌ Error in deep check: {e}")

def quick_check_markers_in_text(text, filename):
    """Quick check if markers are in the extracted text"""
    print(f"\n🔍 QUICK CHECK for 'image[[' in {filename} extracted text")
    
    if not text:
        print("   ❌ No text to check")
        return False
    
    if 'image[[' in text:
        count = text.count('image[[')
        print(f"   ✅ FOUND! {count} 'image[[' markers")
        
        # Find positions
        positions = []
        pos = text.find('image[[')
        while pos != -1 and len(positions) < 5:
            positions.append(pos)
            pos = text.find('image[[', pos + 1)
        
        print(f"   First {len(positions)} positions: {positions}")
        
        # Show first marker
        if positions:
            first_pos = positions[0]
            sample = text[max(0, first_pos-30):first_pos+70]
            print(f"   First marker context: ...{sample}...")
        
        return True
    else:
        print("   ❌ 'image[[' NOT FOUND in extracted text")
        
        # Check for any image-related text
        image_indicators = ['image', 'Image', 'IMG', 'figure', 'Figure', 'FIG']
        for indicator in image_indicators:
            if indicator in text:
                count = text.count(indicator)
                print(f"   ℹ️  Found '{indicator}': {count} times")
        
        return False

# Add this function to check the raw PDF structure
def check_pdf_structure_for_markers(pdf_path):
    """Check PDF structure to understand why markers aren't extractable"""
    print(f"\n🔍 CHECKING PDF STRUCTURE: {os.path.basename(pdf_path)}")
    
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            print(f"   PDF has {len(reader.pages)} pages")
            
            # Check each page for text vs graphics
            for page_num in range(min(5, len(reader.pages))):
                page = reader.pages[page_num]
                text = page.extract_text()
                
                print(f"\n   Page {page_num + 1}:")
                print(f"   - Text length: {len(text)} chars")
                
                # Look for ANY markers or image references
                patterns = [
                    r'image\[\[', r'image', r'Image', r'IMG',
                    r'figure', r'Figure', r'FIG',
                    r'graphic', r'Graphic', r'GRAPHIC'
                ]
                
                found_patterns = []
                import re
                for pattern in patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        found_patterns.append(pattern)
                
                if found_patterns:
                    print(f"   - Found patterns: {', '.join(found_patterns)}")
                    
                    # Show context of first 'figure' mention
                    if 'figure' in [p.lower() for p in found_patterns]:
                        for line in text.split('\n'):
                            if 'figure' in line.lower():
                                print(f"   - Context: {line.strip()[:100]}")
                                break
                else:
                    print(f"   - No image/pattern references found")
                
                # Check if page has XObjects (images/graphics)
                if '/XObject' in page.get('/Resources', {}):
                    print(f"   - Has XObjects (images/graphics): ✓")
                    
                    # Count them
                    xobjects = page['/Resources']['/XObject']
                    count = len(xobjects) if hasattr(xobjects, '__len__') else 1
                    print(f"   - XObject count: {count}")
                    
                    # Check if any are images
                    image_count = 0
                    for obj in xobjects:
                        if xobjects[obj].get('/Subtype') == '/Image':
                            image_count += 1
                    print(f"   - Actual images: {image_count}")
                else:
                    print(f"   - Has XObjects (images/graphics): ✗")
        
        print(f"\n📊 ANALYSIS:")
        print(f"   1. The 'image[[...]]' markers are VISUAL but not TEXT")
        print(f"   2. They're likely embedded as graphics/vector objects")
        print(f"   3. PyPDF2 can't extract graphics as text")
        print(f"   4. Found 'figure' mentions - these might be captions")
        
        # Recommendation
        print(f"\n💡 RECOMMENDATION:")
        print(f"   Option 1: Convert PDF to Word DOCX (preserves images better)")
        print(f"   Option 2: Use OCR if PDF is scanned")
        print(f"   Option 3: Extract images separately and upload them")
        print(f"   Option 4: Manually add image references in the text")
        
    except Exception as e:
        print(f"❌ Error checking PDF structure: {e}")

# Also add this function to create manual image references
def add_manual_image_references(text_content, pdf_filename):
    """Add manual image references based on PDF content"""
    print(f"\n🖼️  Adding manual image references for {pdf_filename}")
    
    # Count approximate image positions based on content
    image_count = 0
    
    # Look for image positions in the text structure
    lines = text_content.split('\n')
    enhanced_lines = []
    
    for i, line in enumerate(lines):
        enhanced_lines.append(line)
        
        # Check if this looks like a place where an image should be
        if ('as seen below' in line.lower() or 
            'see figure' in line.lower() or 
            'see image' in line.lower() or
            'diagram' in line.lower() or
            'illustration' in line.lower()):
            
            image_count += 1
            image_ref = f"\n[Image {image_count} - See original PDF for diagram]"
            enhanced_lines.append(image_ref)
            print(f"   Added image reference {image_count} after line: {line[:50]}...")
    
    if image_count > 0:
        print(f"   Added {image_count} manual image references")
    else:
        print(f"   No obvious image positions found")
    
    return '\n'.join(enhanced_lines)






# ============================
# IMAGE EXTRACTION FUNCTIONS
# ============================

def extract_images_from_docx(doc_path, exam_title):
    """Extract images from Word document and prepare for upload"""
    try:
        document = Document(doc_path)
        images = []
        
        # Counter for image naming
        image_counter = 1
        
        # Iterate through all relationships in the document
        for rel in document.part.rels.values():
            if "image" in rel.reltype:  # Check relationship type
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Create unique filename
                    image_filename = f"{exam_title}_image_{image_counter}.png"
                    
                    # Convert to base64 for easy transmission
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
                    
                    images.append({
                        'filename': image_filename,
                        'data': image_base64,
                        'mimetype': 'image/png'  # Default, can be enhanced
                    })
                    
                    image_counter += 1
                    
                except Exception as e:
                    print(f"⚠️  Could not extract image: {e}")
                    continue
                    
        print(f"📸 Extracted {len(images)} images from Word document")
        return images
        
    except Exception as e:
        print(f"❌ Error extracting images from Word: {e}")
        return []

def extract_tables_from_docx(document):
    """Extract tables with proper structure from Word document"""
    tables_data = []
    
    try:
        for table_index, table in enumerate(document.tables):
            table_rows = []
            
            # Extract headers (first row)
            headers = []
            if table.rows:
                header_cells = table.rows[0].cells
                headers = [cell.text.strip() for cell in header_cells]
            
            # Extract all rows
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if row_cells:  # Only add non-empty rows
                    table_rows.append(row_cells)
            
            if table_rows:
                tables_data.append({
                    'headers': headers,
                    'rows': table_rows,
                    'index': table_index
                })
                
        print(f"📊 Extracted {len(tables_data)} tables from Word document")
        return tables_data
        
    except Exception as e:
        print(f"❌ Error extracting tables: {e}")
        return []

def test_image_marker_detection(pdf_path):
    """Quick test to see if image markers are in the PDF"""
    print(f"\n🔍 QUICK TEST: Checking for image markers in {os.path.basename(pdf_path)}")
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Check first few pages
            for page_num in range(min(3, len(pdf_reader.pages))):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                # Simple check - just look for "image[[" in text
                if 'image[[' in text:
                    print(f"   ✅ Page {page_num + 1}: Contains 'image[[' marker")
                    
                    # Find all occurrences
                    markers = []
                    start = 0
                    while True:
                        start = text.find('image[[', start)
                        if start == -1:
                            break
                        end = text.find(']]', start)
                        if end == -1:
                            break
                        marker = text[start:end+2]
                        markers.append(marker)
                        start = end + 2
                    
                    print(f"     Found {len(markers)} markers:")
                    for marker in markers[:3]:  # Show first 3
                        print(f"       {marker}")
                    
                    if len(markers) > 3:
                        print(f"       ... and {len(markers)-3} more")
                else:
                    print(f"   ❌ Page {page_num + 1}: No 'image[[' markers found")
    
    except Exception as e:
        print(f"   ❌ Test failed: {e}")

def extract_images_from_pdf(pdf_path, exam_title):
    """Extract ACTUAL embedded images from PDF using PyMuPDF"""
    print(f"   🔍 Extracting REAL images from PDF using PyMuPDF...")
    
    images = []
    
    try:
        import fitz  # PyMuPDF
        
        # Open the PDF
        doc = fitz.open(pdf_path)
        total_images_extracted = 0
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Get list of images on this page
            image_list = page.get_images(full=True)
            
            if image_list:
                print(f"   Page {page_num + 1}: Found {len(image_list)} image(s)")
                
                for img_index, img_info in enumerate(image_list):
                    try:
                        # Extract the image
                        xref = img_info[0]  # Image reference number
                        base_image = doc.extract_image(xref)
                        
                        if base_image:
                            image_data = base_image["image"]
                            image_ext = base_image["ext"]
                            image_width = base_image["width"]
                            image_height = base_image["height"]
                            
                            # Create filename
                            image_filename = f"{exam_title}_page{page_num+1}_img{img_index+1}.{image_ext}"
                            
                            # Convert to base64
                            image_base64 = base64.b64encode(image_data).decode('utf-8')
                            
                            # Determine MIME type
                            if image_ext.lower() == "jpg" or image_ext.lower() == "jpeg":
                                mimetype = "image/jpeg"
                            elif image_ext.lower() == "png":
                                mimetype = "image/png"
                            elif image_ext.lower() == "gif":
                                mimetype = "image/gif"
                            else:
                                mimetype = f"image/{image_ext}"
                            
                            images.append({
                                'filename': image_filename,
                                'data': image_base64,
                                'mimetype': mimetype,
                                'width': image_width,
                                'height': image_height,
                                'page': page_num + 1,
                                'index': img_index + 1,
                                'size': len(image_data),
                                'is_real_image': True,
                                'format': image_ext
                            })
                            
                            total_images_extracted += 1
                            print(f"     ✓ Extracted: {image_filename} ({image_width}x{image_height}, {len(image_data)} bytes)")
                            
                    except Exception as img_error:
                        print(f"     ✗ Failed to extract image {img_index + 1}: {img_error}")
                        continue
        
        doc.close()
        
        if total_images_extracted > 0:
            print(f"\n   ✅ SUCCESS: Extracted {total_images_extracted} REAL images from PDF")
            print(f"   ℹ️  These are actual embedded images, not text markers")
        else:
            print(f"\n   ℹ️  No embedded images found in PDF")
            
        return images
        
    except ImportError:
        print("   ❌ PyMuPDF not installed. Run: pip install PyMuPDF")
        print("   ℹ️  Falling back to manual reference detection")
        # Fall back to manual detection
        return extract_images_from_pdf_fallback(pdf_path, exam_title)
    except Exception as e:
        print(f"❌ Error extracting images with PyMuPDF: {e}")
        import traceback
        traceback.print_exc()
        return []


def extract_images_from_pdf_fallback(pdf_path, exam_title):
    """Fallback method if PyMuPDF fails"""
    print(f"   🔍 Using fallback method for image detection...")
    
    images = []
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Count images using PyPDF2
            image_count = 0
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                
                # Check for images in resources
                if '/XObject' in page.get('/Resources', {}):
                    xobjects = page['/Resources']['/XObject']
                    for obj in xobjects:
                        if xobjects[obj].get('/Subtype') == '/Image':
                            image_count += 1
                            images.append({
                                'filename': f"{exam_title}_page{page_num+1}_img{image_count}.txt",
                                'data': '',
                                'mimetype': 'text/plain',
                                'is_placeholder': True,
                                'is_manual_ref': True,
                                'page': page_num + 1,
                                'note': f"Image detected but cannot extract. Check original PDF page {page_num + 1}"
                            })
            
            if image_count > 0:
                print(f"   📊 Detected {image_count} images (cannot extract with PyPDF2)")
            else:
                print(f"   ℹ️  No images detected")
                
        return images
        
    except Exception as e:
        print(f"❌ Error in fallback method: {e}")
        return []







def debug_pdf_images(pdf_path):
    """Debug what's actually in the PDF regarding images"""
    print(f"\n🔍 IMAGE DIAGNOSTICS FOR PDF: {pdf_path}")
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            print(f"   Total pages: {len(pdf_reader.pages)}")
            
            for page_num in range(min(3, len(pdf_reader.pages))):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                # Look for image markers
                markers = re.findall(r'image\[\[.*?\]\]', text)
                
                if markers:
                    print(f"\n   Page {page_num + 1}:")
                    print(f"   Found {len(markers)} image markers")
                    for marker in markers:
                        print(f"     Marker: {marker}")
                        
                        # Check if it's coordinates
                        coords_match = re.search(r'\[\[(\d+),\s*(\d+),\s*(\d+),\s*(\d+)\]\]', marker)
                        if coords_match:
                            print(f"     Coordinates: x1={coords_match.group(1)}, y1={coords_match.group(2)}, x2={coords_match.group(3)}, y2={coords_match.group(4)}")
                
                # Check for actual embedded images (PyPDF2 method)
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            print(f"\n   ⚠️ REAL IMAGE FOUND on page {page_num + 1}!")
                            print(f"     Object: {obj}")
                else:
                    print(f"\n   Page {page_num + 1}: No embedded image objects found")
        
        print(f"\n📊 SUMMARY: Your PDF contains TEXT MARKERS for images, not actual embedded images.")
        print(f"   These markers like 'image[[238, 224, 759, 659]]' are just text.")
        print(f"   The actual images are likely separate files or were lost in PDF creation.")
        
    except Exception as e:
        print(f"❌ Error in diagnostics: {e}")

def extract_tables_from_pdf(pdf_path):
    """Extract tables from PDF (basic implementation)"""
    # PDF table extraction is complex - requires libraries like camelot or tabula
    print("ℹ️  PDF table extraction not implemented - requires additional libraries")
    return []

# ============================
# TEXT EXTRACTION FUNCTIONS
# ============================

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF files - ENHANCED with better processing"""
    try:
        print(f"   Extracting text from PDF using PyPDF2...")
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            all_text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text:
                    # Add page marker
                    all_text += f"\n\n===== Page {page_num + 1} =====\n\n"
                    
                    # Clean up the text
                    cleaned_text = clean_pdf_text(page_text)
                    all_text += cleaned_text
            
            print(f"   Extracted {len(all_text)} characters total")
            
            # Save extracted text to a file for debugging
            debug_file = "pdf_extracted_text_debug.txt"
            with open(debug_file, 'w', encoding='utf-8') as f:
                f.write(all_text)
            print(f"   Saved extracted text to {debug_file}")
            
            return all_text
            
    except Exception as e:
        print(f"❌ Error reading PDF {pdf_path}: {e}")
        import traceback
        traceback.print_exc()
        return ""

def clean_pdf_text(text):
    """Clean and format PDF extracted text"""
    if not text:
        return ""
    
    # Split into lines
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Fix common PDF extraction issues
        # 2. Fix line breaks in the middle of sentences
        if (len(cleaned_lines) > 0 and 
            not cleaned_lines[-1].endswith(('.', '!', '?', ':', ';')) and
            not line[0].isupper() and
            len(cleaned_lines[-1]) > 10):
            # Join with previous line
            cleaned_lines[-1] = cleaned_lines[-1] + ' ' + line
        else:
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def extract_text_from_docx(docx_path):
    """Extract text from Word document"""
    try:
        document = Document(docx_path)
        paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        return '\n\n'.join(paragraphs)
    except Exception as e:
        print(f"❌ Error reading DOCX {docx_path}: {e}")
        return ""

# ============================
# CONTENT PARSING FUNCTIONS
# ============================

def convert_table_to_markdown(table_data):
    """Convert table data to markdown format"""
    if not table_data.get('rows'):
        return ""
    
    markdown_table = ""
    
    # Add headers if available
    if table_data.get('headers'):
        headers = table_data['headers']
        markdown_table += "| " + " | ".join(headers) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    else:
        # Use first row as headers if no specific headers
        first_row = table_data['rows'][0] if table_data['rows'] else []
        if first_row:
            markdown_table += "| " + " | ".join(first_row) + " |\n"
            markdown_table += "| " + " | ".join(["---"] * len(first_row)) + " |\n"
            table_data['rows'] = table_data['rows'][1:]  # Remove first row from data
    
    # Add data rows
    for row in table_data.get('rows', []):
        markdown_table += "| " + " | ".join(row) + " |\n"
    
    return markdown_table

def enhance_content_with_media(text_content, images, tables, exam_title):
    """Enhance text content with images and tables in markdown format"""
    enhanced_content = text_content
    
    # Group images by page
    images_by_page = {}
    for image in images:
        page = image.get('page', 1)
        if page not in images_by_page:
            images_by_page[page] = []
        images_by_page[page].append(image)
    
    # Add REAL images first
    real_images = [img for img in images if img.get('is_real_image')]
    for i, image in enumerate(real_images, 1):
        image_reference = f"\n\n![{exam_title} Diagram {i}]({image['filename']})"
        if image.get('width') and image.get('height'):
            image_reference += f"\n*Diagram {i}: {image['width']}x{image['height']} pixels*"
        enhanced_content += image_reference
    
    # Add placeholder images
    placeholder_images = [img for img in images if img.get('is_placeholder') and not img.get('is_real_image')]
    for i, image in enumerate(placeholder_images, len(real_images) + 1):
        if image.get('is_manual_ref'):
            page_num = image.get('page', '?')
            image_reference = f"\n\n**Diagram {i}:** [See original PDF page {page_num}]"
        else:
            image_reference = f"\n\n![{exam_title} Image {i} - See original PDF](image_placeholder_{i})"
        enhanced_content += image_reference
    
    # Add tables to content
    for i, table in enumerate(tables, 1):
        table_markdown = convert_table_to_markdown(table)
        if table_markdown:
            table_reference = f"\n\n**Table {i}:**\n\n{table_markdown}"
            enhanced_content += table_reference
    
    # Add summary if we have images
    if images:
        real_count = len([img for img in images if img.get('is_real_image')])
        placeholder_count = len(images) - real_count
        
        if real_count > 0:
            summary = f"\n\n---\n**Note:** This document contains {real_count} extracted diagram(s)"
            if placeholder_count > 0:
                summary += f" and {placeholder_count} additional image reference(s)"
            summary += "."
            enhanced_content += summary
    
    return enhanced_content








def extract_reading_content_from_text(text, exam_uuid, source_type="general"):
    """
    Extract structured content from text
    source_type: "pdf" or "docx" or "general"
    """
    sections = []
    
    if not text or not text.strip():
        print(f"⚠️  Warning: Empty text content for {source_type}")
        return [{
            'id': str(uuid.uuid4()),
            'exam_id': exam_uuid,
            'text': '## Content Overview ##',
            'options': ['No text content extracted from document.'],
            'correct_idx': -1
        }]
    
    print(f"   Parsing {len(text)} characters of {source_type} text...")
    
    if source_type == "pdf":
        # Special handling for PDFs
        return parse_pdf_content(text, exam_uuid)
    else:
        # General parsing for DOCX and other formats
        return parse_general_content(text, exam_uuid)

def parse_pdf_content(text, exam_uuid):
    """Parse PDF content - simplified approach"""
    sections = []
    
    # Split by page markers
    pages = re.split(r'===== Page \d+ =====', text)
    
    for page_num, page_content in enumerate(pages, 1):
        if not page_content.strip():
            continue
        
        lines = [line.strip() for line in page_content.split('\n') if line.strip()]
        
        if not lines:
            continue
        
        # Find a title for this page
        title = f"Page {page_num}"
        content_lines = []
        
        for line in lines:
            # Skip empty lines and page numbers
            if not line or line.isdigit():
                continue
            
            # Check if this could be a header (for PDFs)
            if (len(line) < 100 and 
                line[0].isalpha() and 
                (line.isupper() or (line[0].isupper() and line[1:].islower())) and
                not line.endswith('.') and
                not ',' in line):
                
                # This might be a header
                if not title.startswith('##'):
                    title = f"## {line} ##"
                else:
                    # Already have a title, add as content
                    content_lines.append(line)
            else:
                # Regular content
                content_lines.append(line)
        
        # If we didn't find a good header, use first line
        if not title.startswith('##') and content_lines:
            first_line = content_lines[0][:80]
            title = f"## Page {page_num}: {first_line}... ##"
        
        # Join content
        content = '\n\n'.join(content_lines) if content_lines else "Content from PDF page."
        
        sections.append({
            'id': str(uuid.uuid4()),
            'exam_id': exam_uuid,
            'text': title,
            'options': [content],
            'correct_idx': -1
        })
    
    # If no sections were created, create one with all content
    if not sections:
        sections.append({
            'id': str(uuid.uuid4()),
            'exam_id': exam_uuid,
            'text': "## PDF Content ##",
            'options': [text],
            'correct_idx': -1
        })
    
    print(f"   Created {len(sections)} sections from PDF")
    return sections

def parse_general_content(text, exam_uuid):
    """Parse general content (for DOCX files)"""
    sections = []
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    current_section = None
    current_content = []
    
    for line in lines:
        # Check if this is a topic/header (already has ## marks from your manual editing)
        if line.startswith('## ') and line.endswith(' ##'):
            # Save previous section
            if current_section and current_content:
                current_section['options'] = ['\n\n'.join(current_content)]
                sections.append(current_section)
                current_content = []
            
            # Start new section
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': line,  # Already has ## marks
                'options': [],
                'correct_idx': -1
            }
        elif current_section is not None:
            # Add to current section's content
            current_content.append(line)
        else:
            # No section started yet, create one
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': f"## {line[:80]}... ##" if len(line) > 80 else f"## {line} ##",
                'options': [],
                'correct_idx': -1
            }
            current_content = [line]
    
    # Don't forget the last section
    if current_section:
        if current_content:
            current_section['options'] = ['\n\n'.join(current_content)]
        sections.append(current_section)
    
    print(f"   Created {len(sections)} sections from general content")
    return sections

# ============================
# UPLOAD FUNCTIONS
# ============================

def upload_images_to_server(images, exam_id, base_url):
    """Create image references (images are embedded in content, not uploaded separately)"""
    uploaded_image_urls = {}
    
    real_images = [img for img in images if img.get('is_real_image')]
    
    print(f"📤 Images will be embedded in content (no separate upload needed)")
    
    # Just create filename references - images are already embedded as base64
    for image in images:
        uploaded_image_urls[image['filename']] = image['filename']
        
        if image.get('is_real_image'):
            print(f"   📸 Embedded: {image['filename']} ({image.get('width', '?')}x{image.get('height', '?')})")
        elif image.get('is_manual_ref'):
            print(f"   📋 Reference: {image['filename']} (page {image.get('page', '?')})")
        else:
            print(f"   📝 Placeholder: {image['filename']}")
    
    return uploaded_image_urls
# ============================
# PROFESSION SELECTION
# ============================

def get_profession_from_user():
    """Ask user which profession to upload reading materials for"""
    disciplines = {
        '1': ('gp', 'General Practitioner'),
        '2': ('nurse', 'Nurse'),
        '3': ('midwife', 'Midwife'),
        '4': ('lab_tech', 'Lab Technologist'),
        '5': ('physiotherapist', 'Physiotherapist'),
        '6': ('icu_nurse', 'ICU Nurse'),
        '7': ('emergency_nurse', 'Emergency Nurse'),
        '8': ('neonatal_nurse', 'Neonatal Nurse'),
        '9': ('pharmacist', 'Pharmacist')
    }
    
    print("\n🎯 SELECT PROFESSION FOR READING MATERIALS:")
    print("=" * 50)
    print("   1. General Practitioner (gp)")
    print("   2. Nurse (nurse)")
    print("   3. Midwife (midwife)")
    print("   4. Lab Technologist (lab_tech)")
    print("   5. Physiotherapist (physiotherapist)")
    print("   6. ICU Nurse (icu_nurse)")
    print("   7. Emergency Nurse (emergency_nurse)")
    print("   8. Neonatal Nurse (neonatal_nurse)")
    print("   9. Pharmacist (pharmacist)")
    print("=" * 50)
    
    while True:
        choice = input("\nEnter your choice (1-9): ").strip()
        if choice in disciplines:
            discipline_id, discipline_name = disciplines[choice]
            print(f"✅ Selected: {discipline_name} (discipline_id: {discipline_id})")
            return discipline_id, discipline_name
        else:
            print("❌ Invalid choice. Please enter a number between 1-9")

def get_reading_materials_folder_path(discipline_id):
    """Get the appropriate reading materials folder path for each discipline"""
    base_path = r'd:\Thecla\Training Examinations'
    
    folder_mapping = {
        'gp': r'GP\Study Notes',
        'nurse': r'Nurses\Study Notes',
        'midwife': r'Midwives\Reading Materials',
        'lab_tech': r'Lab Technologists\Reading Materials',
        'physiotherapist': r'Physiotherapists\Reading Materials',
        'icu_nurse': r'Specialty Nurses\ICU\Reading Materials',
        'emergency_nurse': r'Specialty Nurses\Emergency\Reading Materials',
        'neonatal_nurse': r'Neonate\Notes',
        'pharmacist': r'Pharmacist\Notes'
    }
    
    folder = folder_mapping.get(discipline_id, 'Reading Materials')
    full_path = os.path.join(base_path, folder)
    
    # Create folder if it doesn't exist
    if not os.path.exists(full_path):
        print(f"📁 Creating folder: {full_path}")
        os.makedirs(full_path, exist_ok=True)
    
    return full_path

# ============================
# FILE MANAGEMENT FUNCTIONS
# ============================

def check_file_exists_on_server(filename, discipline_id, base_url):
    """Check if a file with the same title already exists on the server"""
    try:
        exam_title = os.path.splitext(filename)[0]
        
        # First, try to get all exams for this discipline
        response = requests.get(f"{base_url}/exams")
        if response.status_code == 200:
            all_exams = response.json()
            # Look for exams with same title and discipline
            for exam in all_exams:
                if exam.get('title') == exam_title and exam.get('discipline_id') == discipline_id:
                    return True, exam.get('id')  # Return True and the existing exam ID
        return False, None
    except Exception as e:
        print(f"⚠️  Could not check server for existing files: {e}")
        return False, None

def select_files_to_upload(all_files):
    """Let user choose which files to upload"""
    print(f"\n📄 Found {len(all_files)} document file(s):")
    for i, filename in enumerate(all_files, 1):
        print(f"   {i}. {filename}")
    
    while True:
        print("\n📋 UPLOAD OPTIONS:")
        print("   1. Upload ALL files")
        print("   2. Select specific files")
        
        choice = input("\nEnter your choice (1 or 2): ").strip()
        
        if choice == '1':
            return all_files  # Return all files
        elif choice == '2':
            return select_specific_files(all_files)
        else:
            print("❌ Invalid choice. Please enter 1 or 2")

def select_specific_files(all_files):
    """Let user select specific files from the list"""
    selected_files = []
    
    print("\n🔍 SELECT FILES (enter numbers separated by spaces):")
    for i, filename in enumerate(all_files, 1):
        print(f"   {i}. {filename}")
    
    while True:
        try:
            selections = input("\nEnter file numbers (e.g., 1 3 5): ").strip().split()
            if not selections:
                print("❌ Please select at least one file")
                continue
            
            selected_indices = []
            for sel in selections:
                idx = int(sel) - 1
                if 0 <= idx < len(all_files):
                    selected_indices.append(idx)
                else:
                    print(f"❌ Invalid number: {sel}. Please choose between 1-{len(all_files)}")
            
            if selected_indices:
                selected_files = [all_files[i] for i in selected_indices]
                print(f"\n✅ Selected {len(selected_files)} file(s):")
                for filename in selected_files:
                    print(f"   📄 {filename}")
                return selected_files
            else:
                print("❌ No valid files selected")
                
        except ValueError:
            print("❌ Please enter numbers only")

def confirm_file_replacement(filename):
    """Ask user for confirmation before replacing existing file"""
    while True:
        choice = input(f"🔄 File '{filename}' already exists. Replace it? (y/n): ").strip().lower()
        if choice in ['y', 'yes']:
            return True
        elif choice in ['n', 'no']:
            return False
        else:
            print("❌ Please enter 'y' for yes or 'n' for no")






def embed_images_directly(text_content, images, exam_title):
    """Compress and embed images as base64 data URLs to avoid huge payloads"""
    if not images:
        return text_content
    
    result = text_content
    real_images = [img for img in images if img.get('is_real_image') and img.get('data')]
    
    if not real_images:
        return result
    
    print(f"   Compressing and embedding {len(real_images)} images...")
    
    # Try to use PIL for compression if available
    try:
        from PIL import Image
        PIL_AVAILABLE = True
    except ImportError:
        PIL_AVAILABLE = False
        print("   ⚠️  Pillow not installed. Using original images (may be large).")
        print("   💡 Install: pip install Pillow")
    
    import io
    import base64
    
    images_section = "\n\n---\n## Key Diagrams ##\n\n"
    embedded_count = 0
    
    # Store data URLs for inline embedding
    embedded_images = []
    
    # STRATEGY: Process images for embedding
    for i, image in enumerate(real_images, 1):
        # Limit to 4 images maximum
        if embedded_count >= 4:
            break
            
        try:
            original_data = image['data']
            original_size = len(original_data)
            
            # Skip if image is already too large (> 100KB as base64)
            if original_size > 100000 and len(real_images) > 2:  # 100KB
                print(f"     Skipping {image['filename']}: too large ({original_size//1024}KB)")
                continue
            
            compressed_b64 = original_data
            mimetype = image['mimetype']
            
            # Compress if PIL is available
            if PIL_AVAILABLE:
                try:
                    # Decode base64
                    img_data = base64.b64decode(original_data)
                    img = Image.open(io.BytesIO(img_data))
                    
                    width, height = img.size
                    
                    # Resize if too large
                    max_dimension = 500
                    if max(width, height) > max_dimension:
                        ratio = max_dimension / max(width, height)
                        new_size = (int(width * ratio), int(height * ratio))
                        img = img.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # Convert format for better compression
                    buffer = io.BytesIO()
                    
                    # Use JPEG for photos, PNG for diagrams
                    if img.mode == 'RGB' and width * height > 100000:  # Large image
                        img.save(buffer, format='JPEG', quality=70, optimize=True)
                        mimetype = 'image/jpeg'
                    else:
                        img.save(buffer, format='PNG', optimize=True)
                        mimetype = 'image/png'
                    
                    compressed_data = buffer.getvalue()
                    compressed_b64 = base64.b64encode(compressed_data).decode('utf-8')
                    
                    compression_ratio = (1 - len(compressed_b64)/original_size) * 100
                    if compression_ratio > 10:
                        print(f"     Compressed {image['filename']}: {original_size//1024}KB → {len(compressed_b64)//1024}KB")
                    
                except Exception as comp_error:
                    print(f"     Compression failed for {image['filename']}: {comp_error}")
                    # Use original data
            
            # Create data URL
            data_url = f"data:{mimetype};base64,{compressed_b64}"
            
            # Store for inline embedding
            embedded_images.append({
                'data_url': data_url,
                'width': image.get('width'),
                'height': image.get('height'),
                'page': image.get('page'),
                'index': i,
                'filename': image['filename']
            })
            
            # Add to images section at the end
            images_section += f"**Diagram {i}**"
            if image.get('page'):
                images_section += f" (Page {image['page']})"
            images_section += ":\n\n"
            
            images_section += f"![{exam_title} Diagram {i}]({data_url})\n\n"
            
            # Add dimensions if available
            if image.get('width') and image.get('height'):
                images_section += f"*Size: {image['width']}x{image['height']} pixels*\n\n"
            
            embedded_count += 1
            
        except Exception as e:
            print(f"     ✗ Failed to embed {image.get('filename', f'image{i}')}: {e}")
            continue
    
    # Add ACTUAL IMAGES at logical positions in text
    lines = result.split('\n')
    enhanced_lines = []
    image_ref_counter = 0
    
    for line in lines:
        enhanced_lines.append(line)
        
        # Add ACTUAL IMAGES where they're mentioned (not text references)
        if ('as seen below' in line.lower() or 
            'see figure' in line.lower() or 
            'diagram' in line.lower() or
            'illustration' in line.lower()):
            
            if image_ref_counter < len(embedded_images):
                img_data = embedded_images[image_ref_counter]
                data_url = img_data['data_url']
                
                # Add actual image markdown
                image_markdown = f"\n\n![Diagram {img_data['index']}]({data_url})"
                if img_data.get('width') and img_data.get('height'):
                    image_markdown += f"\n*Diagram {img_data['index']}: {img_data['width']}x{img_data['height']} pixels*"
                
                enhanced_lines.append(image_markdown)
                print(f"     Embedded Diagram {img_data['index']} inline")
                image_ref_counter += 1
    
    result = '\n'.join(enhanced_lines)
    
    # Add note about total images
    if real_images:
        images_section += f"\n**Note:** This PDF contains {len(real_images)} diagram(s). "
        if embedded_count < len(real_images):
            images_section += f"Showing {embedded_count} key diagrams here. "
        images_section += "View original PDF for all illustrations.\n"
    
    # Add the images section at the end
    if embedded_count > 0:
        result += images_section
    
    return result



# ============================
# MAIN UPLOAD SCRIPT
# ============================

BASE_URL = "https://thecla-backend.onrender.com"
API_URL = f'{BASE_URL}/exam/'  # Singular endpoint

# Ask user which profession to upload reading materials for
discipline_id, discipline_name = get_profession_from_user()
folder_path = get_reading_materials_folder_path(discipline_id)

print(f"\n📁 Scanning folder: {folder_path}")
print(f"🎯 Uploading reading materials for: {discipline_name}")

# Check if folder exists
if not os.path.exists(folder_path):
    print(f"❌ Folder not found: {folder_path}")
    print("💡 Please create the folder and add reading material documents, or check the path configuration.")
    exit()

# Get all .docx AND .pdf files in the folder
docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.docx') and not f.startswith('~$')]
pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]

all_files = docx_files + pdf_files
# Get all .docx AND .pdf files in the folder
docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.docx') and not f.startswith('~$')]
pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]

all_files = docx_files + pdf_files

# ADD THIS LINE ↓↓↓
all_files = sort_files_by_chapter(all_files)
# ADD THIS LINE ↑↑↑


if not all_files:
    print(f"❌ No document files found in: {folder_path}")
    print("💡 Please add reading material documents (.docx or .pdf files) to the folder and try again.")
    exit()

# FILE SELECTION PROCESS
files_to_upload = select_files_to_upload(all_files)

print(f"\n🚀 Ready to upload {len(files_to_upload)} file(s)")

# Run a diagnostic on first PDF
for filename in files_to_upload:
    if filename.lower().endswith('.pdf'):
        file_path = os.path.join(folder_path, filename)
        print(f"\n🔍 PRE-PROCESSING DIAGNOSTIC for: {filename}")
        check_pdf_for_hidden_markers(file_path)
        break  # Just check first PDF

uploaded_count = 0
skipped_count = 0

for filename in files_to_upload:
    file_path = os.path.join(folder_path, filename)
    print(f"\n{'='*60}")
    print(f"PROCESSING: {filename}")
    print(f"{'='*60}")
    
    try:
        # First, run debug on the file
        if filename.lower().endswith('.pdf'):
            print(f"🔍 Running PDF debug analysis...")
            pdf_has_text = debug_pdf_content(file_path)
            if not pdf_has_text:
                print(f"⚠️  PDF may not have extractable text. Trying anyway...")
        elif filename.lower().endswith('.docx'):
            print(f"🔍 Running DOCX debug analysis...")
            debug_docx_content(file_path)
        
        # CHECK FOR EXISTING FILE ON SERVER
        file_exists, existing_exam_id = check_file_exists_on_server(filename, discipline_id, BASE_URL)
        
        if file_exists:
            if not confirm_file_replacement(filename):
                print(f"⏭️  Skipping '{filename}' - user chose not to replace")
                skipped_count += 1
                continue
            # If replacing, we'll use the existing exam ID instead of generating a new one
            exam_uuid = existing_exam_id or str(uuid.uuid4())
            replacement_note = " (REPLACING EXISTING)"
        else:
            exam_uuid = str(uuid.uuid4())
            replacement_note = " (NEW)"
        
        exam_title = os.path.splitext(filename)[0]  # Remove extension
        
        
            
            # Handle different file types with MEDIA EXTRACTION
        images = []
        tables = []
        full_text = ""
        
        if filename.lower().endswith('.docx'):
            # Process Word document
            print(f"📝 Processing Word document...")
            document = Document(file_path)
            
            # Extract text
            full_text = extract_text_from_docx(file_path)
            file_type = "Word Document"
            
            # EXTRACT IMAGES AND TABLES
            images = extract_images_from_docx(file_path, exam_title)
            tables = extract_tables_from_docx(document)
            
        elif filename.lower().endswith('.pdf'):
            # Process PDF document
            print(f"📝 Processing PDF document...")
            full_text = extract_text_from_pdf(file_path)
            file_type = "PDF Document"
    
            # ADD STRUCTURE CHECK
            check_pdf_structure_for_markers(file_path)
    
            # ADD MANUAL IMAGE REFERENCES
            full_text = add_manual_image_references(full_text, filename)
    
            # ADD DIAGNOSTICS
            check_debug_file_for_markers()
            quick_check_markers_in_text(full_text, filename)
    
            # EXTRACT REAL IMAGES USING PYMUPDF
            images = extract_images_from_pdf(file_path, exam_title)
            tables = extract_tables_from_pdf(file_path)    
            
        else:
            print(f"❌ Unsupported file type: {filename}")
            skipped_count += 1
            continue
        # Check if we got any text
        if not full_text or not full_text.strip():
            print(f"⚠️  WARNING: No text extracted from: {filename}")
            print(f"   Text length: {len(full_text)}")
            
            # Create minimal placeholder text
            if filename.lower().endswith('.pdf'):
                full_text = f"PDF Document: {exam_title}\n\nThis PDF contains content that requires manual review or OCR processing."
            else:
                full_text = f"Document: {exam_title}\n\nContent extraction may require manual review."
        
        print(f"   Extracted text: {len(full_text)} characters")
        print(f"   Found images: {len(images)}")
        print(f"   Found tables: {len(tables)}")





        
                # ENHANCE CONTENT WITH MEDIA
        if images or tables:
            print(f"🎨 Enhancing content with media...")
            enhanced_text = enhance_content_with_media(full_text, images, tables, exam_title)
        else:
            enhanced_text = full_text
        
        # EMBED IMAGES IN CONTENT (No separate upload)
        if images:
            print(f"📤 Embedding {len(images)} images directly in content...")
            
            # Count real images
            real_images = [img for img in images if img.get('is_real_image')]
            
            if real_images:
                print(f"   Found {len(real_images)} real images to embed")
                
                # Estimate total size
                total_size = sum(len(img.get('data', '')) for img in real_images)
                
                if total_size > 500000:  # If > 500KB total
                    print(f"   ⚠️  Images total ~{total_size//1024}KB - will compress and limit")
                
                enhanced_text = embed_images_directly(enhanced_text, images, exam_title)
            else:
                print(f"   ℹ️  No real images to embed")
            
            # Create simple references for media_info
            image_references = {img['filename']: img['filename'] for img in images}
        else:
            image_references = {}
        
        # USE THE PARSER FOR READING MATERIALS
        print(f"📖 Parsing content into structured sections...")
        
        # Determine source type for parser
        source_type = "pdf" if filename.lower().endswith('.pdf') else "docx"
        questions = extract_reading_content_from_text(enhanced_text, exam_uuid, source_type)
        
        # Validate that we have content
        if not questions:
            print(f"⚠️  No content sections extracted from {filename}")
            # Create a minimal section with the title
            questions = [{
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': f"## {exam_title} ##",
                'options': ["Content extracted from document. Please view original file for complete content."],
                'correct_idx': -1
            }]

 
      




        # Prepare payload with media info
        payload = {
            "id": exam_uuid,
            "title": exam_title,
            "discipline_id": discipline_id,
            "time_limit": 50,
            "source": "singular",
            "questions": questions,
            "media_info": {
                "images_count": len(images),
                "tables_count": len(tables),
                "image_references": image_references
            }
        }
        
        print(f"\n📤 UPLOAD SUMMARY:")
        print(f"   File: {filename}{replacement_note}")
        print(f"   Material Title: {exam_title}")
        print(f"   File Type: {file_type}")
        print(f"   Sections: {len(questions)}")
        print(f"   Images: {len(images)}")
        print(f"   Tables: {len(tables)}")
        print(f"   Discipline: {discipline_name} ({discipline_id})")
        
        # Show first section preview
        if questions and questions[0].get('options'):
            first_content = questions[0]['options'][0]
            preview = first_content[:150].replace('\n', ' ')
            print(f"   First section preview: {preview}...")
        
        # Upload to server
        print(f"\n📤 Uploading to server...")
        response = requests.post(API_URL, json=payload)
        
        if response.status_code == 200:
            print(f"✅ SUCCESS: {exam_title} uploaded successfully!")
            uploaded_count += 1
        else:
            print(f"❌ FAILED: {exam_title} - Status: {response.status_code}")
            print(f"   Error details: {response.text[:200]}...")
            skipped_count += 1
            
    except Exception as e:
        print(f"💥 ERROR processing {filename}: {e}")
        import traceback
        traceback.print_exc()
        skipped_count += 1

print(f"\n{'='*60}")
print(f"🎉 UPLOAD COMPLETED!")
print(f"{'='*60}")
print(f"📊 SUMMARY:")
print(f"   Uploaded: {uploaded_count}")
print(f"   Skipped: {skipped_count}")
print(f"   Total: {uploaded_count + skipped_count}")
print(f"🎯 Discipline: {discipline_name}")
print(f"📁 Processed from: {folder_path}")

# Clean up debug file
debug_file = "pdf_extracted_text_debug.txt"
if os.path.exists(debug_file):
    try:
        os.remove(debug_file)
        print(f"🧹 Cleaned up debug file: {debug_file}")
    except:
        pass

print(f"\n💡 TIP: If PDFs show no content, they may be scanned/image-based.")
print(f"      Consider using OCR tools or converting to Word format first.")
print(f"💡 IMAGE ISSUE: If image markers aren't found, they might be in graphics")
print(f"      rather than text. Use PyMuPDF for better text extraction.")
print(f"{'='*60}")