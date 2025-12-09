import os
import re
import uuid
import base64
from docx import Document
from docx.shared import Inches
import requests
import PyPDF2
from PIL import Image
import io

def extract_images_from_docx(doc_path, exam_title):
    """Extract images from Word document and prepare for upload"""
    try:
        document = Document(doc_path)
        images = []
        
        # Counter for image naming
        image_counter = 1
        
        # Iterate through all relationships in the document
        for rel in document.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Create unique filename
                    image_filename = f"{exam_title}_image_{image_counter}.png"
                    
                    # Convert to base64 for easy transmission
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
                    
                    images.append({
                        'filename': image_filename,
                        'data': image_base64,
                        'mimetype': 'image/png'  # Default, can be enhanced
                    })
                    
                    image_counter += 1
                    
                except Exception as e:
                    print(f"⚠️  Could not extract image: {e}")
                    continue
                    
        print(f"📸 Extracted {len(images)} images from Word document")
        return images
        
    except Exception as e:
        print(f"❌ Error extracting images from Word: {e}")
        return []

def extract_tables_from_docx(document):
    """Extract tables with proper structure from Word document"""
    tables_data = []
    
    try:
        for table_index, table in enumerate(document.tables):
            table_rows = []
            
            # Extract headers (first row)
            headers = []
            if table.rows:
                header_cells = table.rows[0].cells
                headers = [cell.text.strip() for cell in header_cells]
            
            # Extract all rows
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if row_cells:  # Only add non-empty rows
                    table_rows.append(row_cells)
            
            if table_rows:
                tables_data.append({
                    'headers': headers,
                    'rows': table_rows,
                    'index': table_index
                })
                
        print(f"📊 Extracted {len(tables_data)} tables from Word document")
        return tables_data
        
    except Exception as e:
        print(f"❌ Error extracting tables: {e}")
        return []

def extract_images_from_pdf(pdf_path, exam_title):
    """Extract images from PDF (basic implementation)"""
    # Note: PDF image extraction is complex and may require additional libraries
    # For now, return empty list - you can enhance this with PyMuPDF or pdf2image
    print("ℹ️  PDF image extraction not implemented - requires additional libraries")
    return []

def extract_tables_from_pdf(pdf_path):
    """Extract tables from PDF (basic implementation)"""
    # PDF table extraction is complex - requires libraries like camelot or tabula
    print("ℹ️  PDF table extraction not implemented - requires additional libraries")
    return []

def convert_table_to_markdown(table_data):
    """Convert table data to markdown format"""
    if not table_data['rows']:
        return ""
    
    markdown_table = ""
    
    # Add headers if available
    if table_data['headers']:
        markdown_table += "| " + " | ".join(table_data['headers']) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(table_data['headers'])) + " |\n"
    else:
        # Use first row as headers if no specific headers
        first_row = table_data['rows'][0]
        markdown_table += "| " + " | ".join(first_row) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(first_row)) + " |\n"
        table_data['rows'] = table_data['rows'][1:]  # Remove first row from data
    
    # Add data rows
    for row in table_data['rows']:
        markdown_table += "| " + " | ".join(row) + " |\n"
    
    return markdown_table

def enhance_content_with_media(text_content, images, tables, exam_title):
    """Enhance text content with images and tables in markdown format"""
    enhanced_content = text_content
    
    # Add images to content
    for i, image in enumerate(images):
        image_reference = f"\n\n![{exam_title} Image {i+1}]({image['filename']})"
        enhanced_content += image_reference
    
    # Add tables to content
    for i, table in enumerate(tables):
        table_markdown = convert_table_to_markdown(table)
        if table_markdown:
            table_reference = f"\n\n**Table {i+1}:**\n\n{table_markdown}"
            enhanced_content += table_reference
    
    return enhanced_content

def upload_images_to_server(images, exam_id, base_url):
    """Upload extracted images to server"""
    uploaded_image_urls = {}
    
    for image in images:
        try:
            # Prepare payload for image upload
            payload = {
                'exam_id': exam_id,
                'filename': image['filename'],
                'image_data': image['data'],
                'mimetype': image['mimetype']
            }
            
            # Upload image to server
            response = requests.post(f"{base_url}/upload-image", json=payload)
            
            if response.status_code == 200:
                result = response.json()
                uploaded_image_urls[image['filename']] = result.get('image_url', image['filename'])
                print(f"✅ Uploaded image: {image['filename']}")
            else:
                print(f"❌ Failed to upload image {image['filename']}: {response.status_code}")
                # Fallback: use filename as reference
                uploaded_image_urls[image['filename']] = image['filename']
                
        except Exception as e:
            print(f"❌ Error uploading image {image['filename']}: {e}")
            uploaded_image_urls[image['filename']] = image['filename']
    
    return uploaded_image_urls

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF files"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        print(f"❌ Error reading PDF {pdf_path}: {e}")
        return ""

def extract_reading_content_from_text(text, exam_uuid):
    """Extract structured content from text with support for media markers"""
    sections = []
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    current_section = None
    
    for i, line in enumerate(lines):
        # Skip empty lines
        if not line:
            continue
            
        # IMPROVED: Focus on numbered subtopics for TOC items
        is_header = (
            re.match(r'^\d+\.\s+[A-Z]', line) or  # "1. Isotretinoin" - MAIN FIX
            re.match(r'^\d+\.\s+"', line) or       # "1. "Which client..." 
            # Keep some original patterns but make them stricter
            (re.match(r'^\d+\.\s+', line) and len(line) < 100) or
            re.match(r'^[IVX]+\.\s+[A-Z]', line)   # "I. Introduction"
        )
        
        if is_header:
            # Save previous section if exists
            if current_section and (current_section['text'] or current_section['options']):
                sections.append(current_section)
            
            # Start new section
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': line,  # This is the section title/header
                'options': [],  # This will store the content paragraphs
                'correct_idx': -1
            }
        
        # If we're in a section and this line is content
        elif current_section is not None:
            # Check if this line contains image or table markers
            if re.search(r'!\[.*\]\(.*\)', line) or re.search(r'^\|.*\|$', line):
                # It's a media line, add as-is to preserve markdown
                current_section['options'].append(line)
            else:
                # Regular text content
                current_section['options'].append(line)
        
        # If no section started yet, create one with the first line as title
        elif current_section is None:
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': "Content Overview",  # Default title
                'options': [line],  # First line as content
                'correct_idx': -1
            }
    
    # Don't forget the last section
    if current_section and (current_section['text'] or current_section['options']):
        sections.append(current_section)
    
    # If no sections found, treat everything as one section
    if not sections and lines:
        sections.append({
            'id': str(uuid.uuid4()),
            'exam_id': exam_uuid,
            'text': "Reading Material",  # Default title
            'options': lines,  # All lines as content
            'correct_idx': -1
        })
    
    print(f"📖 Extracted {len(sections)} sections from document")
    for i, section in enumerate(sections):
        # Count media elements in this section
        media_count = sum(1 for line in section['options'] 
                         if re.search(r'!\[.*\]\(.*\)', line) or re.search(r'^\|.*\|$', line))
        print(f"  Section {i+1}: '{section['text'][:50]}...' - {len(section['options'])} content lines ({media_count} media elements)")
    
    return sections

# PROFESSION SELECTION FOR ALL 8 DISCIPLINES
def get_profession_from_user():
    """Ask user which profession to upload reading materials for"""
    disciplines = {
        '1': ('gp', 'General Practitioner'),
        '2': ('nurse', 'Nurse'),
        '3': ('midwife', 'Midwife'),
        '4': ('lab_tech', 'Lab Technologist'),
        '5': ('physiotherapist', 'Physiotherapist'),
        '6': ('icu_nurse', 'ICU Nurse'),
        '7': ('emergency_nurse', 'Emergency Nurse'),
        '8': ('neonatal_nurse', 'Neonatal Nurse')
    }
    
    print("\n🎯 SELECT PROFESSION FOR READING MATERIALS:")
    print("=" * 50)
    print("   1. General Practitioner (gp)")
    print("   2. Nurse (nurse)")
    print("   3. Midwife (midwife)")
    print("   4. Lab Technologist (lab_tech)")
    print("   5. Physiotherapist (physiotherapist)")
    print("   6. ICU Nurse (icu_nurse)")
    print("   7. Emergency Nurse (emergency_nurse)")
    print("   8. Neonatal Nurse (neonatal_nurse)")
    print("=" * 50)
    
    while True:
        choice = input("\nEnter your choice (1-8): ").strip()
        if choice in disciplines:
            discipline_id, discipline_name = disciplines[choice]
            print(f"✅ Selected: {discipline_name} (discipline_id: {discipline_id})")
            return discipline_id, discipline_name
        else:
            print("❌ Invalid choice. Please enter a number between 1-8")

# DYNAMIC FOLDER PATHS FOR EACH PROFESSION'S READING MATERIALS
def get_reading_materials_folder_path(discipline_id):
    """Get the appropriate reading materials folder path for each discipline"""
    base_path = r'd:\Thecla\Training Examinations'
    
    folder_mapping = {
        'gp': r'GP\Study Notes',
        'nurse': r'Nurses\Study Notes',
        'midwife': r'Midwives\Reading Materials',
        'lab_tech': r'Lab Technologists\Reading Materials',
        'physiotherapist': r'Physiotherapists\Reading Materials',
        'icu_nurse': r'Specialty Nurses\ICU\Reading Materials',
        'emergency_nurse': r'Specialty Nurses\Emergency\Reading Materials',
        'neonatal_nurse': r'Neonate\Notes'
    }
    
    folder = folder_mapping.get(discipline_id, 'Reading Materials')
    full_path = os.path.join(base_path, folder)
    
    # Create folder if it doesn't exist
    if not os.path.exists(full_path):
        print(f"📁 Creating folder: {full_path}")
        os.makedirs(full_path, exist_ok=True)
    
    return full_path

# CHECK IF FILE ALREADY EXISTS ON SERVER
def check_file_exists_on_server(filename, discipline_id, base_url):
    """Check if a file with the same title already exists on the server"""
    try:
        exam_title = os.path.splitext(filename)[0]
        
        # First, try to get all exams for this discipline
        response = requests.get(f"{base_url}/exams")
        if response.status_code == 200:
            all_exams = response.json()
            # Look for exams with same title and discipline
            for exam in all_exams:
                if exam.get('title') == exam_title and exam.get('discipline_id') == discipline_id:
                    return True, exam.get('id')  # Return True and the existing exam ID
        return False, None
    except Exception as e:
        print(f"⚠️  Could not check server for existing files: {e}")
        return False, None

# FILE SELECTION MENU
def select_files_to_upload(all_files):
    """Let user choose which files to upload"""
    print(f"\n📄 Found {len(all_files)} document file(s):")
    for i, filename in enumerate(all_files, 1):
        print(f"   {i}. {filename}")
    
    while True:
        print("\n📋 UPLOAD OPTIONS:")
        print("   1. Upload ALL files")
        print("   2. Select specific files")
        
        choice = input("\nEnter your choice (1 or 2): ").strip()
        
        if choice == '1':
            return all_files  # Return all files
        elif choice == '2':
            return select_specific_files(all_files)
        else:
            print("❌ Invalid choice. Please enter 1 or 2")

# SPECIFIC FILE SELECTION
def select_specific_files(all_files):
    """Let user select specific files from the list"""
    selected_files = []
    
    print("\n🔍 SELECT FILES (enter numbers separated by spaces):")
    for i, filename in enumerate(all_files, 1):
        print(f"   {i}. {filename}")
    
    while True:
        try:
            selections = input("\nEnter file numbers (e.g., 1 3 5): ").strip().split()
            if not selections:
                print("❌ Please select at least one file")
                continue
            
            selected_indices = []
            for sel in selections:
                idx = int(sel) - 1
                if 0 <= idx < len(all_files):
                    selected_indices.append(idx)
                else:
                    print(f"❌ Invalid number: {sel}. Please choose between 1-{len(all_files)}")
            
            if selected_indices:
                selected_files = [all_files[i] for i in selected_indices]
                print(f"\n✅ Selected {len(selected_files)} file(s):")
                for filename in selected_files:
                    print(f"   📄 {filename}")
                return selected_files
            else:
                print("❌ No valid files selected")
                
        except ValueError:
            print("❌ Please enter numbers only")

# CONFIRM FILE REPLACEMENT
def confirm_file_replacement(filename):
    """Ask user for confirmation before replacing existing file"""
    while True:
        choice = input(f"🔄 File '{filename}' already exists. Replace it? (y/n): ").strip().lower()
        if choice in ['y', 'yes']:
            return True
        elif choice in ['n', 'no']:
            return False
        else:
            print("❌ Please enter 'y' for yes or 'n' for no")

# MAIN UPLOAD SCRIPT FOR READING MATERIALS
#BASE_URL = "https://93c2af60c5f9.ngrok-free.app"
BASE_URL = "https://thecla-backend.onrender.com"
API_URL = f'{BASE_URL}/exam/'  # Singular endpoint

# Ask user which profession to upload reading materials for
discipline_id, discipline_name = get_profession_from_user()
folder_path = get_reading_materials_folder_path(discipline_id)

print(f"\n📁 Scanning folder: {folder_path}")
print(f"🎯 Uploading reading materials for: {discipline_name}")

# Check if folder exists
if not os.path.exists(folder_path):
    print(f"❌ Folder not found: {folder_path}")
    print("💡 Please create the folder and add reading material documents, or check the path configuration.")
    exit()

# Get all .docx AND .pdf files in the folder
docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.docx') and not f.startswith('~$')]
pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]

all_files = docx_files + pdf_files

if not all_files:
    print(f"❌ No document files found in: {folder_path}")
    print("💡 Please add reading material documents (.docx or .pdf files) to the folder and try again.")
    exit()

# FILE SELECTION PROCESS
files_to_upload = select_files_to_upload(all_files)

print(f"\n🚀 Ready to upload {len(files_to_upload)} file(s)")

uploaded_count = 0
skipped_count = 0

for filename in files_to_upload:
    file_path = os.path.join(folder_path, filename)
    try:
        # CHECK FOR EXISTING FILE ON SERVER
        file_exists, existing_exam_id = check_file_exists_on_server(filename, discipline_id, BASE_URL)
        
        if file_exists:
            if not confirm_file_replacement(filename):
                print(f"⏭️  Skipping '{filename}' - user chose not to replace")
                skipped_count += 1
                continue
            # If replacing, we'll use the existing exam ID instead of generating a new one
            exam_uuid = existing_exam_id or str(uuid.uuid4())
            replacement_note = " (REPLACING EXISTING)"
        else:
            exam_uuid = str(uuid.uuid4())
            replacement_note = " (NEW)"
        
        exam_title = os.path.splitext(filename)[0]  # Remove extension
        
        # Handle different file types with MEDIA EXTRACTION
        images = []
        tables = []
        full_text = ""
        
        if filename.lower().endswith('.docx'):
            # Process Word document with MEDIA EXTRACTION
            document = Document(file_path)
            
            # Extract text
            full_text = '\n'.join([para.text for para in document.paragraphs])
            file_type = "Word Document"
            
            # EXTRACT IMAGES AND TABLES
            images = extract_images_from_docx(file_path, exam_title)
            tables = extract_tables_from_docx(document)
            
        elif filename.lower().endswith('.pdf'):
            # Process PDF document
            full_text = extract_text_from_pdf(file_path)
            file_type = "PDF Document"
            
            # PDF media extraction (basic - can be enhanced)
            images = extract_images_from_pdf(file_path, exam_title)
            tables = extract_tables_from_pdf(file_path)
            
        else:
            print(f"❌ Unsupported file type: {filename}")
            skipped_count += 1
            continue
        
        if not full_text.strip():
            print(f"⚠️  No text extracted from: {filename}")
            skipped_count += 1
            continue
        
        # ENHANCE CONTENT WITH MEDIA
        if images or tables:
            print(f"🎨 Enhancing content with {len(images)} images and {len(tables)} tables")
            enhanced_text = enhance_content_with_media(full_text, images, tables, exam_title)
        else:
            enhanced_text = full_text
        
        # UPLOAD IMAGES TO SERVER FIRST
        image_references = {}
        if images:
            print(f"📤 Uploading {len(images)} images to server...")
            image_references = upload_images_to_server(images, exam_uuid, BASE_URL)
        
        # USE THE PARSER FOR READING MATERIALS (with enhanced content)
        questions = extract_reading_content_from_text(enhanced_text, exam_uuid)
        
        # Prepare payload with media info
        payload = {
            "id": exam_uuid,
            "title": exam_title,
            "discipline_id": discipline_id,
            "time_limit": 50,
            "source": "singular",
            "questions": questions,
            "media_info": {
                "images_count": len(images),
                "tables_count": len(tables),
                "image_references": image_references
            }
        }
        
        print(f"\n📤 Uploading reading material from file: {file_path}{replacement_note}")
        print(f"📝 Material Title: {exam_title}")
        print(f"📄 File Type: {file_type}")
        print(f"📖 Sections: {len(questions)}")
        print(f"🖼️  Images: {len(images)}")
        print(f"📊 Tables: {len(tables)}")
        print(f"🎯 Discipline: {discipline_name} ({discipline_id})")
        
        response = requests.post(API_URL, json=payload)
        if response.status_code == 200:
            print(f"✅ SUCCESS: {exam_title} - Status: {response.status_code}")
            uploaded_count += 1
        else:
            print(f"❌ FAILED: {exam_title} - Status: {response.status_code}")
            print(f"   Error: {response.text}")
            skipped_count += 1
            
    except Exception as e:
        print(f"💥 Error processing {file_path}: {e}")
        skipped_count += 1

print(f"\n🎉 Upload completed!")
print(f"📊 Summary: {uploaded_count} uploaded, {skipped_count} skipped")
print(f"🎯 Discipline: {discipline_name}")
print(f"📁 Processed from: {folder_path}")