import os
from docx import Document
import re

def full_document_analysis(docx_path):
    print(f"\n{'='*60}")
    print(f"FULL DOCUMENT ANALYSIS")
    print(f"{'='*60}")
    
    doc = Document(docx_path)
    all_paragraphs = doc.paragraphs
    
    print(f"\n📊 STATISTICS:")
    print(f"   Total paragraphs: {len(all_paragraphs)}")
    
    # Show EVERY paragraph with its content
    print(f"\n📄 COMPLETE PARAGRAPH LIST:")
    print(f"   {'Idx':<5} {'Content Preview'}")
    print(f"   {'-'*5} {'-'*50}")
    
    for i, para in enumerate(all_paragraphs):
        text = para.text.strip()
        if text:
            preview = text[:80].replace('\n', ' ')
            print(f"   [{i:3}] {preview}")
        else:
            print(f"   [{i:3}] [EMPTY PARAGRAPH]")
    
    # Check for specific content from pages 13-17
    print(f"\n🔍 SEARCHING FOR PAGES 13-17 CONTENT:")
    
    search_terms = [
        "Septic Shock",
        "SVT",
        "Bradycardia", 
        "Supraventricular Tachycardia",
        "Cardiogenic and Septic Shock",
        "Warm Shock",
        "Cold Shock"
    ]
    
    found_content = {}
    for term in search_terms:
        for i, para in enumerate(all_paragraphs):
            if term.lower() in para.text.lower():
                if term not in found_content:
                    found_content[term] = []
                found_content[term].append(i)
    
    if found_content:
        print(f"   ✅ Found clinical content:")
        for term, indices in found_content.items():
            print(f"      - '{term}' at paragraphs: {indices}")
    else:
        print(f"   ❌ NONE of the clinical content found!")
        print(f"   This means pages 13-17 content is NOT in the paragraphs!")
    
    # Check if content might be in tables
    print(f"\n📊 TABLES:")
    print(f"   Number of tables: {len(doc.tables)}")
    if doc.tables:
        for t_idx, table in enumerate(doc.tables):
            print(f"   Table {t_idx + 1}: {len(table.rows)} rows x {len(table.columns)} cols")
            for r_idx, row in enumerate(table.rows[:3]):  # First 3 rows
                row_text = ' | '.join([cell.text.strip()[:30] for cell in row.cells])
                if row_text:
                    print(f"      Row {r_idx}: {row_text}")
    
    # Check headers and footers
    print(f"\n📌 HEADERS & FOOTERS:")
    for s_idx, section in enumerate(doc.sections):
        print(f"   Section {s_idx + 1}:")
        if section.header:
            header_text = ' '.join([p.text for p in section.header.paragraphs if p.text.strip()])
            if header_text:
                print(f"      Header: {header_text[:100]}")
        if section.footer:
            footer_text = ' '.join([p.text for p in section.footer.paragraphs if p.text.strip()])
            if footer_text:
                print(f"      Footer: {footer_text[:100]}")
    
    # Save all text to file for inspection
    with open("document_full_text.txt", "w", encoding="utf-8") as f:
        for i, para in enumerate(all_paragraphs):
            f.write(f"[{i}] {para.text}\n")
    
    print(f"\n💾 Full text saved to: document_full_text.txt")
    
    # Recommendation
    print(f"\n🎯 CONCLUSION:")
    if not found_content:
        print(f"   ❌ The clinical content (Septic Shock, SVT, Bradycardia) is NOT in the main paragraphs!")
        print(f"   Possible causes:")
        print(f"      1. Content is in text boxes (python-docx doesn't extract these)")
        print(f"      2. Content is in floating shapes")
        print(f"      3. Content is in comments or revisions")
        print(f"      4. The Word document has corruption")
        print(f"\n   💡 SOLUTION: Copy the content from the Word document and paste into a NEW document")
    else:
        print(f"   ✅ Content IS in the paragraphs at indices: {list(found_content.values())}")
        print(f"   Then the issue must be in your parse_general_content function")
        print(f"   Share the output above for further debugging")

# Run it
file_path = r"d:\Thecla\Training Examinations\Neonate\Notes\check.docx"
full_document_analysis(file_path)