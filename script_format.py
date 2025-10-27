import os
import re
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class ExamDocumentStandardizer:
    def __init__(self):
        self.setup_patterns()
    
    def setup_patterns(self):
        """Setup more flexible regex patterns for parsing questions"""
        self.question_pattern = re.compile(r'^(\d+\.)\s*(.+)$')
        # More flexible option patterns
        self.option_patterns = [
            re.compile(r'^([A-D])\.\s*(.+)$'),  # A. Option text
            re.compile(r'^([A-D])\s+(.+)$'),    # A Option text
            re.compile(r'^\(([A-D])\)\s*(.+)$'), # (A) Option text
            re.compile(r'^([a-d])\.\s*(.+)$'),  # a. Option text (lowercase)
        ]
        self.answer_pattern = re.compile(r'^[✅]?\s*(?:Correct\s*)?Answer:\s*([A-Da-d])\.?\s*["]?(.+?)["]?$', re.IGNORECASE)
        self.rationale_pattern = re.compile(r'^Rationale:\s*(.+)$', re.IGNORECASE)
    
    def extract_formatting(self, paragraph):
        """Extract formatting information from a paragraph"""
        formatting = {
            'font_name': None,
            'font_size': None,
            'bold': False,
            'italic': False,
            'underline': False,
            'color': None,
            'alignment': paragraph.alignment,
            'line_spacing': paragraph.paragraph_format.line_spacing,
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after
        }
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            formatting['bold'] = first_run.bold
            formatting['italic'] = first_run.italic
            formatting['underline'] = first_run.underline
            
            if first_run.font.name:
                formatting['font_name'] = first_run.font.name
            if first_run.font.size:
                formatting['font_size'] = first_run.font.size
            if first_run.font.color.rgb:
                formatting['color'] = first_run.font.color.rgb
        
        return formatting
    
    def apply_formatting(self, paragraph, formatting, force_bold=False):
        """Apply formatting to a paragraph with optional forced bold"""
        if not paragraph.runs:
            return
        
        # Apply paragraph-level formatting
        paragraph.alignment = formatting.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if formatting.get('line_spacing'):
            paragraph.paragraph_format.line_spacing = formatting['line_spacing']
        if formatting.get('space_before'):
            paragraph.paragraph_format.space_before = formatting['space_before']
        if formatting.get('space_after'):
            paragraph.paragraph_format.space_after = formatting['space_after']
        
        # Apply run-level formatting
        for run in paragraph.runs:
            if formatting.get('font_name'):
                run.font.name = formatting['font_name']
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), formatting['font_name'])
            
            if formatting.get('font_size'):
                run.font.size = formatting['font_size']
            
            # Force bold for questions, otherwise use original formatting
            if force_bold:
                run.bold = True
            else:
                run.bold = formatting.get('bold', False)
            
            run.italic = formatting.get('italic', False)
            run.underline = formatting.get('underline', False)
            
            if formatting.get('color'):
                run.font.color.rgb = formatting['color']
    
    def is_option_line(self, text):
        """Check if text matches any option pattern"""
        for pattern in self.option_patterns:
            match = pattern.match(text)
            if match:
                return match.groups()
        return None
    
    def parse_question_blocks(self, doc):
        """Parse document into structured question blocks"""
        questions = []
        current_question = None
        collecting_options = False
        option_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            formatting = self.extract_formatting(paragraph)
            
            # Check if this is a new question
            question_match = self.question_pattern.match(text)
            if question_match:
                # Save previous question if exists
                if current_question and current_question.get('question_text'):
                    questions.append(current_question)
                
                # Start new question
                current_question = {
                    'question_number': question_match.group(1),
                    'question_text': question_match.group(2),
                    'question_format': formatting,
                    'options': {},
                    'option_formats': {},
                    'answer_letter': '',
                    'answer_text': '',
                    'answer_format': None,
                    'rationale': '',
                    'rationale_format': None
                }
                collecting_options = True
                option_count = 0
                continue
            
            # If we're not in a question yet, skip
            if not current_question:
                continue
            
            # Check for options with multiple patterns
            option_match = self.is_option_line(text)
            if option_match and collecting_options and option_count < 4:
                letter = option_match[0].upper()
                option_text = option_match[1]
                current_question['options'][letter] = option_text
                current_question['option_formats'][letter] = formatting
                option_count += 1
                continue
            
            # Check for answer
            answer_match = self.answer_pattern.match(text)
            if answer_match:
                collecting_options = False
                current_question['answer_letter'] = answer_match.group(1).upper()
                current_question['answer_text'] = answer_match.group(2)
                current_question['answer_format'] = formatting
                continue
            
            # Check for rationale
            rationale_match = self.rationale_pattern.match(text)
            if rationale_match:
                current_question['rationale'] = rationale_match.group(1)
                current_question['rationale_format'] = formatting
                continue
            
            # If we have fewer than 4 options and we're still collecting, this might be an option in different format
            if collecting_options and option_count < 4 and len(text) < 100:  # Assume options are short
                # Try to parse as option without letter (A, B, C, D order)
                letters = ['A', 'B', 'C', 'D']
                if option_count < len(letters):
                    letter = letters[option_count]
                    current_question['options'][letter] = text
                    current_question['option_formats'][letter] = formatting
                    option_count += 1
                    print(f"    ⚠️ Assigned option {letter} to: {text[:50]}...")
        
        # Process the last question
        if current_question and current_question.get('question_text'):
            questions.append(current_question)
        
        return questions
    
    def debug_document_structure(self, doc, max_lines=50):
        """Debug function to see document structure"""
        print("=== DOCUMENT STRUCTURE ===")
        lines_shown = 0
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Check what type of line this is
                line_type = "Other"
                if self.question_pattern.match(text):
                    line_type = "QUESTION"
                elif self.is_option_line(text):
                    line_type = "OPTION"
                elif self.answer_pattern.match(text):
                    line_type = "ANSWER"
                elif self.rationale_pattern.match(text):
                    line_type = "RATIONALE"
                
                print(f"Line {i:3d} [{line_type:8}]: '{text}'")
                lines_shown += 1
                if lines_shown >= max_lines:
                    print("... (showing first 50 lines)")
                    break
        print("=== END STRUCTURE ===")
    
    def create_standardized_document(self, questions: List[Dict]) -> Document:
        """Create a new document with standardized format"""
        new_doc = Document()
        
        for i, question in enumerate(questions, 1):
            print(f"  📝 Processing question {i}: {question['question_text'][:50]}...")
            
            # 1. QUESTION - FORCE BOLD
            question_para = new_doc.add_paragraph()
            question_para.add_run(f"{question['question_number']} {question['question_text']}")
            self.apply_formatting(question_para, question['question_format'], force_bold=True)
            
            # 2. OPTIONS - Ensure A, B, C, D order
            option_letters = ['A', 'B', 'C', 'D']
            options_found = []
            
            for letter in option_letters:
                if letter in question['options']:
                    option_para = new_doc.add_paragraph()
                    option_text = question['options'][letter]
                    option_para.add_run(f"{letter}. {option_text}")
                    
                    if letter in question['option_formats']:
                        self.apply_formatting(option_para, question['option_formats'][letter], force_bold=False)
                    else:
                        self.apply_formatting(option_para, question['question_format'], force_bold=False)
                    
                    options_found.append(letter)
            
            if not options_found:
                print(f"    ⚠️ No options found for question {i}")
            
            # Add blank line
            new_doc.add_paragraph()
            
            # 3. ANSWER
            if question['answer_letter'] and question['answer_text']:
                answer_para = new_doc.add_paragraph()
                answer_para.add_run(f"Answer: {question['answer_letter']}. {question['answer_text']}")
                if question['answer_format']:
                    self.apply_formatting(answer_para, question['answer_format'], force_bold=False)
            else:
                print(f"    ⚠️ Missing answer for question {i}")
            
            # 4. RATIONALE
            if question['rationale']:
                rationale_para = new_doc.add_paragraph()
                rationale_para.add_run(f"Rationale: {question['rationale']}")
                if question['rationale_format']:
                    self.apply_formatting(rationale_para, question['rationale_format'], force_bold=False)
            else:
                print(f"    ⚠️ Missing rationale for question {i}")
            
            # Add separation between questions
            new_doc.add_paragraph()
        
        return new_doc
    
    def process_document(self, input_path: Path, output_path: Path):
        """Process a single Word document"""
        print(f"📄 Processing: {input_path.name}")
        
        # Skip temporary Word files
        if input_path.name.startswith('~$'):
            print(f"  ⏭️ Skipping temporary file: {input_path.name}")
            return
        
        try:
            # Read the document
            doc = Document(input_path)
            
            # Debug: show document structure with line types
            self.debug_document_structure(doc)
            
            # Parse questions with formatting
            questions = self.parse_question_blocks(doc)
            
            if not questions:
                print(f"  ❌ No questions found in {input_path.name}")
                return
            
            print(f"  ✅ Found {len(questions)} questions")
            
            # Show detailed question info
            for i, q in enumerate(questions, 1):
                print(f"    Q{i}: {q['question_text'][:50]}...")
                print(f"       Options found: {list(q['options'].keys())}")
                for letter, text in q['options'].items():
                    print(f"         {letter}. {text[:40]}...")
                print(f"       Answer: {q['answer_letter']} - {q['answer_text'][:30]}...")
                print(f"       Rationale: {'Yes' if q['rationale'] else 'No'}")
            
            # Create standardized document
            new_doc = self.create_standardized_document(questions)
            
            # Save the new document
            new_doc.save(output_path)
            print(f"  💾 Saved to: {output_path.name}")
            
        except Exception as e:
            print(f"  ❌ Error: {str(e)}")

def main():
    # Your specific paths
    INPUT_DIR = Path(r"d:\Thecla\Training Examinations\Nurses\Prometric Exam\Rationale")
    OUTPUT_DIR = Path(r"d:\Thecla\Training Examinations\Nurses\Prometric Exam\Standardized_Rationale")
    
    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    # Initialize the standardizer
    standardizer = ExamDocumentStandardizer()
    
    # Process all .docx files (skip temp files)
    docx_files = [f for f in INPUT_DIR.glob('*.docx') if not f.name.startswith('~$')]
    
    if not docx_files:
        print(f"❌ No Word documents found in {INPUT_DIR}")
        return
    
    print(f"🔍 Found {len(docx_files)} documents to process")
    print("=" * 60)
    
    for input_file in docx_files:
        output_file = OUTPUT_DIR / f"standardized_{input_file.name}"
        standardizer.process_document(input_file, output_file)
        print("-" * 50)
    
    print("🎉 Standardization complete!")

if __name__ == "__main__":
    main()