import os
import re
import uuid
import base64
from docx import Document
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
import requests
import PyPDF2
from PIL import Image
import io
import json
from datetime import datetime
import zipfile

# Try to import optional dependencies with fallbacks
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF not installed. PDF image extraction will be limited.")
    print("   Install with: pip install PyMuPDF")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("⚠️ pdf2image not installed. Scanned PDF extraction will be limited.")
    print("   Install with: pip install pdf2image")

try:
    import tabula
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False
    print("⚠️ tabula-py not installed. PDF table extraction will be limited.")
    print("   Install with: pip install tabula-py")

# ============================
# DEBUG FUNCTIONS
# ============================

def debug_pdf_content(pdf_path):
    """Debug what's actually in the PDF"""
    print(f"\n🔍 DEBUGGING PDF: {pdf_path}")
    
    if not os.path.exists(pdf_path):
        print(f"❌ File not found: {pdf_path}")
        return False
    
    file_size = os.path.getsize(pdf_path)
    print(f"   File size: {file_size} bytes ({file_size/1024:.1f} KB)")
    
    try:
        with open(pdf_path, 'rb') as f:
            first_bytes = f.read(20)
            print(f"   File signature (hex): {first_bytes.hex()[:40]}...")
            
            if b'%PDF' in first_bytes:
                print(f"   ✓ Valid PDF signature found")
                pdf_version = first_bytes[5:8].decode('ascii', errors='ignore')
                print(f"   PDF version: {pdf_version}")
            else:
                print(f"   ✗ WARNING: Not a valid PDF file")
                return False
    except Exception as e:
        print(f"   File read error: {e}")
        return False
    
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            print(f"\n   PyPDF2 Analysis:")
            print(f"   - Pages found: {len(reader.pages)}")
            print(f"   - Is encrypted: {reader.is_encrypted}")
            
            pages_with_text = 0
            total_text_chars = 0
            
            for i in range(min(3, len(reader.pages))):
                text = reader.pages[i].extract_text()
                text_length = len(text.strip())
                total_text_chars += text_length
                
                if text_length > 0:
                    pages_with_text += 1
                    print(f"\n   Page {i+1}: Has text: ✓ ({text_length} characters)")
                    print(f"   Preview: {text[:200]}...")
                else:
                    print(f"\n   Page {i+1}: Has text: ✗ (0 characters)")
            
            print(f"\n   Summary: Pages with text: {pages_with_text}/{min(3, len(reader.pages))}")
            
            if total_text_chars == 0:
                print("   ⚠️  CRITICAL: No text extracted! PDF is likely scanned/image-based")
                if PDF2IMAGE_AVAILABLE:
                    print("   ℹ️  Will attempt OCR-based extraction")
                return False
            else:
                print("   ✓ Good amount of text extracted")
                return True
                
    except Exception as e:
        print(f"   PyPDF2 error: {e}")
        return False

def debug_docx_content(docx_path):
    """Debug what's in a Word document"""
    print(f"\n🔍 DEBUGGING DOCX: {docx_path}")
    
    try:
        document = Document(docx_path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        print(f"   Paragraphs with text: {len(paragraphs)}")
        
        for i, para in enumerate(paragraphs[:5]):
            preview = para[:100] + "..." if len(para) > 100 else para
            print(f"   Para {i+1}: '{preview}'")
        
        images_count = 0
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                images_count += 1
        print(f"   Images found: {images_count}")
        print(f"   Tables found: {len(document.tables)}")
        
        return True
        
    except Exception as e:
        print(f"   DOCX error: {e}")
        return False

# ============================
# IMAGE EXTRACTION FUNCTIONS
# ============================

def extract_images_from_docx(doc_path, exam_title):
    """Extract ALL images from Word document by unzipping the DOCX file"""
    images = []
    image_counter = 1
    
    try:
        # A .docx file is just a zip archive - extract all images directly
        with zipfile.ZipFile(doc_path, 'r') as docx_zip:
            # Look for image files in the word/media folder
            for file_info in docx_zip.filelist:
                filename = file_info.filename
                
                # Check if it's an image file
                if filename.startswith('word/media/') and any(ext in filename.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']):
                    try:
                        # Read the image data
                        image_data = docx_zip.read(filename)
                        
                        # Get file extension
                        file_ext = os.path.splitext(filename)[1].lower().replace('.', '')
                        if file_ext == 'jpg':
                            file_ext = 'jpeg'
                        
                        # Create filename and encode
                        image_filename = f"{exam_title}_image_{image_counter}.{file_ext}"
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        images.append({
                            'filename': image_filename,
                            'data': image_base64,
                            'mimetype': f'image/{file_ext}',
                            'size': len(image_data)
                        })
                        
                        print(f"   ✓ Extracted: {os.path.basename(filename)}")
                        image_counter += 1
                        
                    except Exception as e:
                        print(f"   ⚠️ Could not extract {filename}: {e}")
                        continue
        
        print(f"📸 Extracted {len(images)} images from Word document")
        return images
        
    except Exception as e:
        print(f"❌ Error extracting images from Word: {e}")
        return []

def extract_images_from_pdf(pdf_path, exam_title):
    """Extract images from PDF using multiple methods"""
    images = []
    
    if not PYMUPDF_AVAILABLE and not PDF2IMAGE_AVAILABLE:
        print("   ⚠️ No PDF image extraction libraries available")
        return []
    
    # Method 1: Using PyMuPDF for embedded images
    if PYMUPDF_AVAILABLE:
        try:
            print("   Attempting to extract embedded images with PyMuPDF...")
            doc = fitz.open(pdf_path)
            image_counter = 1
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_data = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        if image_ext.lower() not in ['png', 'jpg', 'jpeg', 'gif']:
                            image_ext = 'png'
                        
                        image_filename = f"{exam_title}_image_{image_counter}.{image_ext}"
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        images.append({
                            'filename': image_filename,
                            'data': image_base64,
                            'mimetype': f'image/{image_ext}',
                            'page': page_num + 1,
                            'size': len(image_data)
                        })
                        image_counter += 1
                    except Exception as e:
                        print(f"   ⚠️ Could not extract image {img_index}: {e}")
                        continue
            
            doc.close()
            if images:
                print(f"   ✓ Extracted {len(images)} embedded images from PDF")
        except Exception as e:
            print(f"   ⚠️ PyMuPDF extraction failed: {e}")
    
    # Method 2: Using pdf2image for scanned PDFs (if no images found)
    if not images and PDF2IMAGE_AVAILABLE:
        try:
            print("   No embedded images found, trying to render PDF pages as images...")
            # Limit to first 5 pages to avoid huge files
            pdf_images = convert_from_path(pdf_path, dpi=100)
            
            for page_num, img in enumerate(pdf_images[:5], 1):
                # Compress image
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG', optimize=True, quality=85)
                image_data = img_byte_arr.getvalue()
                
                image_filename = f"{exam_title}_page_{page_num}.png"
                image_base64 = base64.b64encode(image_data).decode('utf-8')
                
                images.append({
                    'filename': image_filename,
                    'data': image_base64,
                    'mimetype': 'image/png',
                    'page': page_num,
                    'size': len(image_data)
                })
            
            if images:
                print(f"   ✓ Extracted {len(images)} page images from PDF")
        except Exception as e:
            print(f"   ⚠️ pdf2image extraction failed: {e}")
    
    print(f"   📸 Total images extracted from PDF: {len(images)}")
    return images

# ============================
# TABLE EXTRACTION FUNCTIONS
# ============================

def extract_tables_from_docx(document):
    """Extract tables from Word document"""
    tables_data = []
    
    try:
        for table_index, table in enumerate(document.tables):
            table_rows = []
            headers = []
            
            # Extract headers from first row if available
            if table.rows and len(table.rows) > 0:
                header_cells = table.rows[0].cells
                headers = [clean_text(cell.text) for cell in header_cells]
            
            # Extract all rows
            for row in table.rows:
                row_cells = [clean_text(cell.text) for cell in row.cells]
                if any(row_cells):  # Only add if not empty
                    table_rows.append(row_cells)
            
            if table_rows:
                tables_data.append({
                    'headers': headers,
                    'rows': table_rows,
                    'index': table_index,
                    'rows_count': len(table_rows),
                    'cols_count': len(table_rows[0]) if table_rows else 0
                })
                
        print(f"📊 Extracted {len(tables_data)} tables from Word document")
        return tables_data
        
    except Exception as e:
        print(f"❌ Error extracting tables from DOCX: {e}")
        return []

def extract_tables_from_pdf(pdf_path):
    """Extract tables from PDF using multiple methods"""
    tables_data = []
    
    # Method 1: Using tabula (best for structured tables)
    if TABULA_AVAILABLE:
        try:
            print("   Attempting to extract tables with tabula...")
            tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
            
            for table_idx, table in enumerate(tables):
                if not table.empty:
                    # Clean the table data
                    headers = [clean_text(str(h)) for h in table.columns.tolist()]
                    rows = []
                    
                    for _, row in table.iterrows():
                        clean_row = [clean_text(str(cell)) for cell in row.tolist()]
                        if any(clean_row):
                            rows.append(clean_row)
                    
                    if rows:
                        tables_data.append({
                            'headers': headers,
                            'rows': rows,
                            'index': table_idx,
                            'rows_count': len(rows),
                            'cols_count': len(headers)
                        })
            
            if tables_data:
                print(f"   ✓ Extracted {len(tables_data)} tables using tabula")
        except Exception as e:
            print(f"   ⚠️ Tabula extraction failed: {e}")
    
    # Method 2: Try text-based table detection if tabula failed
    if not tables_data and PYMUPDF_AVAILABLE:
        try:
            print("   Attempting text-based table detection...")
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                # Look for table-like patterns (lines with multiple pipe/space separators)
                lines = text.split('\n')
                potential_tables = []
                current_table = []
                
                for line in lines:
                    # Check if line has multiple columns (simplified detection)
                    if '|' in line or (line.count('  ') > 2):
                        current_table.append(line)
                    elif current_table and len(current_table) > 1:
                        potential_tables.append(current_table)
                        current_table = []
                
                if current_table and len(current_table) > 1:
                    potential_tables.append(current_table)
                
                for table_idx, table_lines in enumerate(potential_tables):
                    rows = []
                    for line in table_lines:
                        if '|' in line:
                            cells = [clean_text(cell.strip()) for cell in line.split('|') if cell.strip()]
                        else:
                            cells = [clean_text(cell.strip()) for cell in line.split() if cell.strip()]
                        
                        if cells:
                            rows.append(cells)
                    
                    if rows:
                        tables_data.append({
                            'headers': rows[0] if rows else [],
                            'rows': rows[1:] if len(rows) > 1 else rows,
                            'index': len(tables_data),
                            'rows_count': len(rows),
                            'cols_count': len(rows[0]) if rows else 0
                        })
            
            doc.close()
            if tables_data:
                print(f"   ✓ Extracted {len(tables_data)} tables using text detection")
                
        except Exception as e:
            print(f"   ⚠️ Text-based table detection failed: {e}")
    
    print(f"   📊 Total tables extracted from PDF: {len(tables_data)}")
    return tables_data

# ============================
# TEXT EXTRACTION FUNCTIONS
# ============================

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF files with better error handling"""
    all_text = ""
    
    try:
        print(f"   Extracting text from PDF using PyPDF2...")
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        all_text += f"\n\n===== Page {page_num + 1} =====\n\n"
                        all_text += page_text
                    else:
                        print(f"   ⚠️ Page {page_num + 1}: No text extracted")
                        
                except Exception as e:
                    print(f"   ⚠️ Error extracting page {page_num + 1}: {e}")
                    continue
            
            print(f"   Extracted {len(all_text)} characters total")
            
            if not all_text.strip() and PDF2IMAGE_AVAILABLE:
                print("   No text found, attempting OCR with pdf2image...")
                # Note: You would need pytesseract for actual OCR
                all_text = "PDF appears to be scanned. Text extraction limited."
            
            return all_text
            
    except Exception as e:
        print(f"❌ Error reading PDF {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Extract text from Word document preserving structure"""
    try:
        document = Document(docx_path)
        text_blocks = []
        
        # Extract paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                text_blocks.append(para.text.strip())
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_blocks.append(row_text)
        
        return '\n\n'.join(text_blocks)
        
    except Exception as e:
        print(f"❌ Error reading DOCX {docx_path}: {e}")
        return ""

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    # Remove extra whitespace
    text = ' '.join(text.split())
    # Remove non-printable characters
    text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
    return text.strip()

# ============================
# CONTENT PARSING FUNCTIONS
# ============================

def convert_table_to_markdown(table_data):
    """Convert table data to markdown format"""
    if not table_data.get('rows'):
        return ""
    
    markdown_table = []
    
    # Determine headers
    if table_data.get('headers') and any(table_data['headers']):
        headers = table_data['headers']
        markdown_table.append("| " + " | ".join(headers) + " |")
        markdown_table.append("| " + " | ".join(["---"] * len(headers)) + " |")
        start_row = 0
    else:
        # Use first row as headers
        first_row = table_data['rows'][0] if table_data['rows'] else []
        if first_row:
            markdown_table.append("| " + " | ".join(first_row) + " |")
            markdown_table.append("| " + " | ".join(["---"] * len(first_row)) + " |")
            start_row = 1
        else:
            start_row = 0
    
    # Add data rows
    for row in table_data.get('rows', [])[start_row:]:
        markdown_table.append("| " + " | ".join(row) + " |")
    
    return '\n'.join(markdown_table)

def enhance_content_with_media(text_content, images, tables, exam_title):
    """Replace image and table markers with actual content inline"""
    
    enhanced_content = text_content
    
    # ========== HANDLE IMAGES ==========
    if images:
        print(f"   🖼️ Looking for image markers...")
        
        # Sort images by number
        def get_image_number(img):
            import re
            match = re.search(r'image[_\-]?(\d+)', img.get('filename', ''), re.IGNORECASE)
            return int(match.group(1)) if match else 999
        
        sorted_images = sorted(images, key=get_image_number)
        
        for i, img in enumerate(sorted_images, 1):
            img_markdown = f"\n\n![{img['filename']}](data:{img['mimetype']};base64,{img['data']})\n\n"
            
            patterns = [f'Image {i}', f'image {i}', f'IMAGE {i}', f'[Image {i}]', f'{{Image {i}}}']
            
            replaced = False
            for pattern in patterns:
                if pattern in enhanced_content:
                    enhanced_content = enhanced_content.replace(pattern, img_markdown)
                    print(f"      ✓ Replaced '{pattern}' with image {i}")
                    replaced = True
                    break
            
            if not replaced:
                import re
                if re.search(rf'(?i)Image\s*{i}', enhanced_content):
                    enhanced_content = re.sub(rf'(?i)Image\s*{i}', img_markdown, enhanced_content, count=1)
                    print(f"      ✓ Replaced 'Image {i}' with image {i}")
                    replaced = True
            
            if not replaced:
                enhanced_content += img_markdown
                print(f"      ⚠️ No marker 'Image {i}' found, appended at end")
    
    # ========== HANDLE TABLES ==========
    if tables:
        print(f"   📊 Looking for table markers...")
        
        for i, table in enumerate(tables, 1):
            table_markdown = convert_table_to_markdown(table)
            if table_markdown:
                # Wrap table in markdown formatting
                table_content = f"\n\n**Table {i}:**\n\n{table_markdown}\n\n"
                
                # Patterns to look for table markers
                patterns = [f'Table {i}', f'table {i}', f'TABLE {i}', f'[Table {i}]', f'{{Table {i}}}']
                
                replaced = False
                for pattern in patterns:
                    if pattern in enhanced_content:
                        enhanced_content = enhanced_content.replace(pattern, table_content)
                        print(f"      ✓ Replaced '{pattern}' with table {i}")
                        replaced = True
                        break
                
                if not replaced:
                    import re
                    if re.search(rf'(?i)Table\s*{i}', enhanced_content):
                        enhanced_content = re.sub(rf'(?i)Table\s*{i}', table_content, enhanced_content, count=1)
                        print(f"      ✓ Replaced 'Table {i}' with table {i}")
                        replaced = True
                
                if not replaced:
                    enhanced_content += table_content
                    print(f"      ⚠️ No marker 'Table {i}' found, appended at end")
    
    return enhanced_content


def parse_content_into_sections(text, exam_uuid, source_type="general"):
    """Parse content into structured sections"""
    sections = []
    
    if not text or not text.strip():
        return [{
            'id': str(uuid.uuid4()),
            'exam_id': exam_uuid,
            'text': 'Content Overview',
            'options': ['No content extracted from document.'],
            'correct_idx': -1
        }]
    
    # Split by common section markers
    lines = text.split('\n')
    current_section = None
    current_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line looks like a heading
        is_heading = False
        
        # Heading indicators
        if len(line) < 100 and line[0].isalpha() and line[0].isupper():
            # Check if it's not a complete sentence
            if not line.endswith('.') and not line.endswith('?') and not line.endswith('!'):
                # Check if it's short enough to be a heading
                if len(line) < 80:
                    # Check if it's not part of a paragraph (not followed by typical content markers)
                    is_heading = True
        
        if is_heading and current_section is None:
            # Start new section
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': line,
                'options': [],
                'correct_idx': -1
            }
        elif is_heading and current_section:
            # Save previous section and start new one
            if current_content:
                current_section['options'] = ['\n\n'.join(current_content)]
                sections.append(current_section)
            
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': line,
                'options': [],
                'correct_idx': -1
            }
            current_content = []
        elif current_section:
            current_content.append(line)
        else:
            # Create a default section
            current_section = {
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': line[:80] + "..." if len(line) > 80 else line,
                'options': [],
                'correct_idx': -1
            }
            current_content = [line]
    
    # Don't forget the last section
    if current_section:
        if current_content:
            current_section['options'] = ['\n\n'.join(current_content)]
        else:
            current_section['options'] = ['Content from document.']
        sections.append(current_section)
    
    # If no sections were created, create a default one
    if not sections:
        sections.append({
            'id': str(uuid.uuid4()),
            'exam_id': exam_uuid,
            'text': 'Document Content',
            'options': [text[:500] + "..." if len(text) > 500 else text],
            'correct_idx': -1
        })
    
    print(f"   Created {len(sections)} sections from content")
    return sections

# ============================
# PROFESSION SELECTION
# ============================

def get_profession_from_user():
    """Ask user which profession to upload reading materials for"""
    disciplines = {
        '1': ('gp', 'General Practitioner'),
        '2': ('nurse', 'Nurse'),
        '3': ('midwife', 'Midwife'),
        '4': ('lab_tech', 'Lab Technologist'),
        '5': ('physiotherapist', 'Physiotherapist'),
        '6': ('icu_nurse', 'ICU Nurse'),
        '7': ('emergency_nurse', 'Emergency Nurse'),
        '8': ('neonatal_nurse', 'Neonatal Nurse'),
        '9': ('pharmacist', 'Pharmacist')
    }
    
    print("\n🎯 SELECT PROFESSION FOR READING MATERIALS:")
    print("=" * 50)
    for key, value in disciplines.items():
        print(f"   {key}. {value[1]} ({value[0]})")
    print("=" * 50)
    
    while True:
        choice = input("\nEnter your choice (1-9): ").strip()
        if choice in disciplines:
            discipline_id, discipline_name = disciplines[choice]
            print(f"✅ Selected: {discipline_name} (discipline_id: {discipline_id})")
            return discipline_id, discipline_name
        else:
            print("❌ Invalid choice. Please enter a number between 1-9")

def get_reading_materials_folder_path(discipline_id):
    """Get the appropriate reading materials folder path for each discipline"""
    base_path = r'd:\Thecla\Training Examinations'
    
    folder_mapping = {
        'gp': r'GP\Study Notes',
        'nurse': r'Nurses\Study Notes',
        'midwife': r'Midwives\Reading Materials',
        'lab_tech': r'Lab Technologists\Reading Materials',
        'physiotherapist': r'Physiotherapists\Reading Materials',
        'icu_nurse': r'Specialty Nurses\ICU\Reading Materials',
        'emergency_nurse': r'Specialty Nurses\Emergency\Reading Materials',
        'neonatal_nurse': r'Neonate\Notes',
        'pharmacist': r'Pharmacist\Notes'
    }
    
    folder = folder_mapping.get(discipline_id, 'Reading Materials')
    full_path = os.path.join(base_path, folder)
    
    if not os.path.exists(full_path):
        print(f"📁 Creating folder: {full_path}")
        os.makedirs(full_path, exist_ok=True)
    
    return full_path

# ============================
# FILE MANAGEMENT FUNCTIONS
# ============================

def check_file_exists_on_server(filename, discipline_id, base_url):
    """Check if a file with the same title already exists on the server"""
    try:
        exam_title = os.path.splitext(filename)[0]
        response = requests.get(f"{base_url}/exams", timeout=10)
        if response.status_code == 200:
            all_exams = response.json()
            for exam in all_exams:
                if exam.get('title') == exam_title and exam.get('discipline_id') == discipline_id:
                    return True, exam.get('id')
        return False, None
    except Exception as e:
        print(f"⚠️  Could not check server for existing files: {e}")
        return False, None

def select_files_to_upload(all_files, folder_path):
    """Let user choose which files to upload"""
    print(f"\n📄 Found {len(all_files)} document file(s):")
    for i, filename in enumerate(all_files, 1):
        file_size = os.path.getsize(os.path.join(folder_path, filename))
        size_kb = file_size / 1024
        print(f"   {i}. {filename} ({size_kb:.1f} KB)")
    
    while True:
        print("\n📋 UPLOAD OPTIONS:")
        print("   1. Upload ALL files")
        print("   2. Select specific files")
        
        choice = input("\nEnter your choice (1 or 2): ").strip()
        
        if choice == '1':
            return all_files
        elif choice == '2':
            return select_specific_files(all_files)
        else:
            print("❌ Invalid choice. Please enter 1 or 2")

def select_specific_files(all_files):
    """Let user select specific files from the list"""
    selected_files = []
    
    print("\n🔍 SELECT FILES (enter numbers separated by spaces):")
    for i, filename in enumerate(all_files, 1):
        print(f"   {i}. {filename}")
    
    while True:
        try:
            selections = input("\nEnter file numbers (e.g., 1 3 5): ").strip().split()
            if not selections:
                print("❌ Please select at least one file")
                continue
            
            selected_indices = []
            for sel in selections:
                idx = int(sel) - 1
                if 0 <= idx < len(all_files):
                    selected_indices.append(idx)
                else:
                    print(f"❌ Invalid number: {sel}")
            
            if selected_indices:
                selected_files = [all_files[i] for i in selected_indices]
                print(f"\n✅ Selected {len(selected_files)} file(s):")
                for filename in selected_files:
                    print(f"   📄 {filename}")
                return selected_files
            else:
                print("❌ No valid files selected")
                
        except ValueError:
            print("❌ Please enter numbers only")

def confirm_file_replacement(filename):
    """Ask user for confirmation before replacing existing file"""
    while True:
        choice = input(f"🔄 File '{filename}' already exists. Replace it? (y/n): ").strip().lower()
        if choice in ['y', 'yes']:
            return True
        elif choice in ['n', 'no']:
            return False
        else:
            print("❌ Please enter 'y' for yes or 'n' for no")

# ============================
# MAIN UPLOAD FUNCTION
# ============================

def upload_single_file(file_path, filename, discipline_id, discipline_name, base_url):
    """Upload a single file to the server"""
    print(f"\n{'='*60}")
    print(f"PROCESSING: {filename}")
    print(f"{'='*60}")
    
    try:
        # Debug the file first
        if filename.lower().endswith('.pdf'):
            debug_pdf_content(file_path)
        elif filename.lower().endswith('.docx'):
            debug_docx_content(file_path)
        
        # Check if file already exists
        file_exists, existing_exam_id = check_file_exists_on_server(filename, discipline_id, base_url)
        
        if file_exists:
            if not confirm_file_replacement(filename):
                print(f"⏭️  Skipping '{filename}' - user chose not to replace")
                return False, "skipped"
            exam_uuid = existing_exam_id or str(uuid.uuid4())
            replacement_note = " (REPLACING EXISTING)"
        else:
            exam_uuid = str(uuid.uuid4())
            replacement_note = " (NEW)"
        
        exam_title = os.path.splitext(filename)[0]
        
        # Extract content based on file type
        images = []
        tables = []
        full_text = ""
        
        if filename.lower().endswith('.docx'):
            print(f"📝 Processing Word document...")
            document = Document(file_path)
            full_text = extract_text_from_docx(file_path)
            file_type = "Word Document"
            images = extract_images_from_docx(file_path, exam_title)
            tables = extract_tables_from_docx(document)
            
        elif filename.lower().endswith('.pdf'):
            print(f"📝 Processing PDF document...")
            full_text = extract_text_from_pdf(file_path)
            file_type = "PDF Document"
            images = extract_images_from_pdf(file_path, exam_title)
            tables = extract_tables_from_pdf(file_path)
        
        # Handle case with no text extracted
        if not full_text or not full_text.strip():
            print(f"⚠️  WARNING: No text extracted from: {filename}")
            if filename.lower().endswith('.pdf'):
                full_text = f"PDF Document: {exam_title}\n\nThis PDF appears to be scanned. Please refer to the original document for content."
            else:
                full_text = f"Document: {exam_title}\n\nContent extraction may require manual review of the original file."
        
        print(f"   Extracted text: {len(full_text)} characters")
        print(f"   Found images: {len(images)}")
        print(f"   Found tables: {len(tables)}")
        
        # Enhance content with tables and images
        if images or tables:
            print(f"🎨 Enhancing content with media...")
            enhanced_text = enhance_content_with_media(full_text, images, tables, exam_title)


            
        else:
            enhanced_text = full_text
        
        # Parse into structured sections
        print(f"📖 Parsing content into structured sections...")
        source_type = "pdf" if filename.lower().endswith('.pdf') else "docx"
        sections = parse_content_into_sections(enhanced_text, exam_uuid, source_type)
        
        if not sections:
            print(f"⚠️  No content sections extracted from {filename}")
            sections = [{
                'id': str(uuid.uuid4()),
                'exam_id': exam_uuid,
                'text': exam_title,
                'options': ["Content extracted from document. Please view original file for complete content."],
                'correct_idx': -1
            }]
        
        # Print upload summary
        print(f"\n📤 UPLOAD SUMMARY:")
        print(f"   File: {filename}{replacement_note}")
        print(f"   Material Title: {exam_title}")
        print(f"   File Type: {file_type}")
        print(f"   Sections: {len(sections)}")
        print(f"   Images: {len(images)}")
        print(f"   Tables: {len(tables)}")
        print(f"   Discipline: {discipline_name} ({discipline_id})")
        
        if sections and sections[0].get('options'):
            first_content = sections[0]['options'][0][:150].replace('\n', ' ')
            print(f"   First section preview: {first_content}...")
        
        # Debug: Check if images are in payload
        if images:
            print(f"\n   🖼️ VERIFYING IMAGES IN PAYLOAD:")
            for i, img in enumerate(images[:3], 1):
                print(f"      Image {i}: {img['filename']}")
                print(f"      - Size: {len(img['data'])} chars base64")
                print(f"      - Preview: {img['data'][:50]}...")
        else:
            print(f"\n   ⚠️ NO IMAGES in payload!")
        
        # Prepare payload
        media_data = {
            "images_count": len(images),
            "tables_count": len(tables),
            "images": images,  # Include full image data
            "tables": tables   # Include table data
        }
        
        payload = {
            "id": exam_uuid,
            "title": exam_title,
            "discipline_id": discipline_id,
            "time_limit": 50,
            "source": "reading_material",
            "questions": sections,
            "media_info": media_data
        }
        
        print(f"\n📤 Uploading to server...")
        response = requests.post(f"{base_url}/exam/", json=payload, timeout=30)
        
        if response.status_code == 200:
            print(f"✅ SUCCESS: {exam_title} uploaded successfully!")
            return True, "success"
        else:
            print(f"❌ FAILED: {exam_title} - Status: {response.status_code}")
            print(f"   Error details: {response.text[:200]}...")
            return False, "failed"
            
    except Exception as e:
        print(f"💥 ERROR processing {filename}: {e}")
        import traceback
        traceback.print_exc()
        return False, "error"

# ============================
# MAIN SCRIPT
# ============================

if __name__ == "__main__":
    # Configuration
    BASE_URL = "https://thecla-backend.onrender.com"
    #BASE_URL = "https://7c8d-94-207-206-21.ngrok-free.app"
    
    # Print banner
    print("\n" + "="*60)
    print("📚 READING MATERIALS UPLOADER")
    print("="*60)
    print(f"Backend URL: {BASE_URL}")
    print(f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Check dependencies
    print("\n🔧 DEPENDENCY CHECK:")
    print(f"   PyMuPDF: {'✅' if PYMUPDF_AVAILABLE else '❌'} (for PDF images)")
    print(f"   pdf2image: {'✅' if PDF2IMAGE_AVAILABLE else '❌'} (for scanned PDFs)")
    print(f"   tabula-py: {'✅' if TABULA_AVAILABLE else '❌'} (for PDF tables)")
    
    # Get profession and folder
    discipline_id, discipline_name = get_profession_from_user()
    folder_path = get_reading_materials_folder_path(discipline_id)
    
    print(f"\n📁 Scanning folder: {folder_path}")
    print(f"🎯 Uploading reading materials for: {discipline_name}")
    
    if not os.path.exists(folder_path):
        print(f"❌ Folder not found: {folder_path}")
        exit()
    
    # Find all documents
    docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.docx') and not f.startswith('~$')]
    pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]
    
    all_files = docx_files + pdf_files
    
    if not all_files:
        print(f"❌ No document files found in: {folder_path}")
        print(f"   Supported formats: .docx, .pdf")
        exit()
    
    # Select files to upload
    # Select files to upload
    files_to_upload = select_files_to_upload(all_files, folder_path)  # ← MUST have folder_path
    
    print(f"\n🚀 Ready to upload {len(files_to_upload)} file(s)")
    
    # Upload files
    uploaded_count = 0
    skipped_count = 0
    failed_count = 0
    
    for filename in files_to_upload:
        file_path = os.path.join(folder_path, filename)
        success, status = upload_single_file(file_path, filename, discipline_id, discipline_name, BASE_URL)
        
        if status == "success":
            uploaded_count += 1
        elif status == "skipped":
            skipped_count += 1
        else:
            failed_count += 1
    
    # Print final summary
    print(f"\n{'='*60}")
    print(f"🎉 UPLOAD COMPLETED!")
    print(f"{'='*60}")
    print(f"📊 SUMMARY:")
    print(f"   ✅ Uploaded: {uploaded_count}")
    print(f"   ⏭️  Skipped: {skipped_count}")
    print(f"   ❌ Failed: {failed_count}")
    print(f"   📁 Total: {uploaded_count + skipped_count + failed_count}")
    print(f"🎯 Discipline: {discipline_name}")
    print(f"📁 Processed from: {folder_path}")
    print(f"{'='*60}")